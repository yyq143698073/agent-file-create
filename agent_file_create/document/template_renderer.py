import hashlib
import json
import logging
import re
from pathlib import Path
from typing import Dict, List, Optional, Set

logger = logging.getLogger(__name__)


# ── Section hash cache ──────────────────────────────────────────────────────

def _section_hashes(content_dict: Dict[str, str]) -> Dict[str, str]:
    """Return md5 hex for each section in *content_dict*."""
    return {
        k: hashlib.md5((v or "").encode()).hexdigest()
        for k, v in (content_dict or {}).items()
    }


def _load_render_cache(output_dir: str) -> Dict[str, str]:
    """Load previous section→hash map from *output_dir/_render_cache.json*."""
    p = Path(output_dir) / "_render_cache.json"
    try:
        if p.exists():
            data = json.loads(p.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                return {str(k): str(v) for k, v in data.items()}
    except Exception:
        pass
    return {}


def _save_render_cache(output_dir: str, hashes: Dict[str, str]) -> None:
    """Persist section→hash map."""
    p = Path(output_dir) / "_render_cache.json"
    try:
        p.write_text(json.dumps(hashes, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception as e:
        logger.warning(f"render_cache_write_failed err={str(e)[:160]}")


def _changed_sections(
    content_dict: Dict[str, str], output_dir: str
) -> Set[str]:
    """Return set of section keys whose content changed since last render."""
    current = _section_hashes(content_dict)
    cached = _load_render_cache(output_dir)
    changed: set[str] = set()
    for k, h in current.items():
        if cached.get(k) != h:
            changed.add(k)
    for k in cached:
        if k not in current:
            changed.add(k)
    return changed


# ── Template placeholder scanning ───────────────────────────────────────────

def _scan_md_placeholders(template_path: str) -> Set[str]:
    """Extract placeholder names from a markdown template."""
    try:
        text = Path(template_path).read_text(encoding="utf-8")
        return set(re.findall(r"\{\{(.+?)\}\}", text))
    except Exception:
        return set()


def _scan_docx_placeholders(template_path: str) -> Set[str]:
    """Extract placeholder names from a DOCX template."""
    keys: set[str] = set()
    try:
        from docx import Document
        doc = Document(template_path)
        for para in doc.paragraphs:
            keys.update(re.findall(r"\{\{(.+?)\}\}", para.text))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        keys.update(re.findall(r"\{\{(.+?)\}\}", para.text))
    except Exception:
        pass
    return keys


from functools import lru_cache


@lru_cache(maxsize=64)
def _get_template_placeholders(template_path: str) -> Set[str]:
    """Return cached set of placeholder names referenced by *template_path*."""
    tp = str(template_path)
    suf = Path(tp).suffix.lower()
    if suf == ".docx":
        return _scan_docx_placeholders(tp)
    elif suf == ".md":
        return _scan_md_placeholders(tp)
    else:
        return set()


def should_skip_render(
    template_path: str, changed: Set[str], *, is_pdf: bool = False
) -> bool:
    """Return True if *template_path* does NOT reference any *changed* section.

    PDF templates are programmatic (reportlab draws all sections), so they are
    only skipped when *changed* is completely empty.
    """
    if not changed:
        return True
    if is_pdf:
        return False  # PDF always references all sections
    refs = _get_template_placeholders(template_path)
    if not refs:
        # Template has no detectable placeholders — render anyway (safety)
        return False
    return not bool(refs & changed)


def build_content_dict(markdown_content: str, template_keys: list = None) -> Dict[str, str]:
    """Parse markdown content into a dict keyed by ``##`` heading text.

    If *template_keys* is provided, any key that does not have an exact
    heading match will be resolved via fuzzy matching against the actual
    headings in the document (using SequenceMatcher).
    """
    from difflib import SequenceMatcher

    lines = (markdown_content or "").splitlines()
    out: Dict[str, str] = {}
    cur = None
    buf: list[str] = []
    for line in lines:
        m2 = re.match(r"^\s*##\s+(.+?)\s*$", line)
        if m2:
            if cur is not None:
                out[cur] = "\n".join(buf).strip()
            cur = m2.group(1).strip()
            buf = []
            continue
        if cur is not None:
            buf.append(line)
    if cur is not None:
        out[cur] = "\n".join(buf).strip()

    for k in list(out.keys()):
        if out.get(k):
            continue
        start = f"## {k}"
        block: list[str] = []
        in_block = False
        for line in lines:
            if line.strip() == start:
                in_block = True
                continue
            if in_block and re.match(r"^\s*##\s+", line):
                break
            if in_block:
                if line.strip().startswith("### "):
                    block.append(line.strip())
                elif block:
                    block.append(line)
        if block:
            out[k] = "\n".join(block).strip()

    title = ""
    for line in lines:
        m1 = re.match(r"^\s*#\s+(.+?)\s*$", line)
        if m1:
            title = m1.group(1).strip()
            break
    if title:
        out["title"] = title

    # ── Fuzzy match: map template keys → nearest actual heading ──
    match_info = {"exact": [], "fuzzy": [], "unmatched": []}
    if template_keys:
        system_keys = {"title", "task_id", "document_outline", "document_content"}
        headings = [k for k in out if k not in system_keys]
        for tk in template_keys:
            if tk in system_keys:
                continue
            if tk in out and out[tk]:
                # Already has content from exact match
                match_info["exact"].append(tk)
                continue
            if tk in out:
                # Key exists but content is empty — try fuzzy
                pass
            if headings:
                best_h, best_s = None, 0.0
                for h in headings:
                    s = SequenceMatcher(None, tk, h).ratio()
                    if tk in h or h in tk:
                        s += 0.3
                    if s > best_s:
                        best_s, best_h = s, h
                if best_h and best_s >= 0.35:
                    out[tk] = out[best_h]
                    match_info["fuzzy"].append(f"{tk} → {best_h}")
                    logger.debug("build_content_dict fuzzy: %r → %r (score=%.2f)", tk, best_h, best_s)
                else:
                    match_info["unmatched"].append(tk)
            else:
                match_info["unmatched"].append(tk)

    out["_template_match_info"] = match_info
    return out


def render_markdown_template(template_path: str, content_dict: Dict[str, str], output_path: str) -> None:
    tp = Path(template_path)
    rendered = tp.read_text(encoding="utf-8")
    for k, v in (content_dict or {}).items():
        rendered = rendered.replace("{{" + str(k) + "}}", str(v))
    Path(output_path).write_text(rendered, encoding="utf-8")


def render_word_template(template_path: str, content_dict: Dict[str, str], output_path: str) -> None:
    try:
        from docxtpl import DocxTemplate

        doc = DocxTemplate(template_path)
        doc.render(content_dict or {})
        doc.save(output_path)
        return
    except Exception:
        pass

    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError("缺少 docxtpl 或 python-docx 依赖，无法生成 docx") from e

    doc = Document(template_path)
    mapping = {str(k): str(v) for k, v in (content_dict or {}).items()}
    for para in doc.paragraphs:
        for k, v in mapping.items():
            token = "{{" + k + "}}"
            if token in para.text:
                for run in para.runs:
                    if token in run.text:
                        run.text = run.text.replace(token, v)
    doc.save(output_path)


def render_pdf_template(template_path: str, content_dict: Dict[str, str], output_path: str) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfgen import canvas
    except Exception as e:
        raise RuntimeError("缺少 reportlab 依赖，无法生成 pdf") from e

    c = canvas.Canvas(output_path, pagesize=A4)
    w, h = A4
    x = 46
    y = h - 60

    font_name = "Helvetica"
    try:
        base_dir = Path(__file__).resolve().parent.parent.parent
        simhei = base_dir / "resource" / "SimHei.ttf"
        if simhei.exists():
            pdfmetrics.registerFont(TTFont("SimHei", str(simhei)))
            font_name = "SimHei"
    except Exception:
        font_name = "Helvetica"

    c.setFont(font_name, 16)
    title = str((content_dict or {}).get("title") or "").strip()
    if title:
        c.drawString(x, y, title)
        y -= 26

    c.setFont(font_name, 11)
    for k, v in (content_dict or {}).items():
        if k == "title":
            continue
        txt = str(v or "").strip()
        if not txt:
            continue
        head = str(k).strip()
        c.setFont(font_name, 12)
        c.drawString(x, y, head)
        y -= 18
        c.setFont(font_name, 11)
        for line in txt.splitlines():
            line = line.strip()
            if not line:
                y -= 10
                continue
            if y <= 60:
                c.showPage()
                y = h - 60
                c.setFont(font_name, 11)
            c.drawString(x, y, line[:120])
            y -= 14
        y -= 10

    c.save()


def render_template(template_path: str, content_dict: Dict[str, str], output_path: str) -> None:
    suf = Path(template_path).suffix.lower().lstrip(".")
    if suf == "md":
        render_markdown_template(template_path, content_dict, output_path)
    elif suf == "docx":
        render_word_template(template_path, content_dict, output_path)
    elif suf == "pdf":
        render_pdf_template(template_path, content_dict, output_path)
    else:
        render_markdown_template(template_path, content_dict, output_path)


def _render_markdown_to_docx(content: str, outline: str, output_path: str) -> None:
    """Generate a DOCX directly from markdown content (no template needed).

    Converts headings (# ## ###), lists (- * 1.), tables (|...|), and
    inline formatting (**bold** *italic*) to python-docx elements.
    """
    from docx import Document

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.font.size = 140000  # 14pt

    lines = (content or "").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        i += 1

        if line.startswith("# ") and len(line) > 2:
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## ") and len(line) > 3:
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### ") and len(line) > 4:
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_run(p, line[2:])
        elif re.match(r"^\d+[.)]\s", line):
            p = doc.add_paragraph(style="List Number")
            _add_formatted_run(p, re.sub(r"^\d+[.)]\s*", "", line))
        elif line.startswith("|") and line.endswith("|"):
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            if rows:
                data = [r for r in rows if not all(re.match(r"^[-:]+$", c) for c in r if c)]
                if data:
                    t = doc.add_table(rows=len(data), cols=max(len(r) for r in data), style="Table Grid")
                    for ri, rd in enumerate(data):
                        for ci, ct in enumerate(rd):
                            if ci < max(len(r) for r in data):
                                t.cell(ri, ci).text = ct
                                if ri == 0:
                                    for pa in t.cell(ri, ci).paragraphs:
                                        for ru in pa.runs:
                                            ru.bold = True
        elif not line:
            continue
        else:
            p = doc.add_paragraph()
            _add_formatted_run(p, line)

    doc.save(output_path)


def _add_formatted_run(para, text: str) -> None:
    """Add text with **bold** and *italic* markers to a paragraph."""
    import re as _re
    segments = _re.split(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*[^*\n]+\*)", text)
    for seg in segments:
        if seg.startswith("***") and seg.endswith("***"):
            run = para.add_run(seg[3:-3])
            run.bold = True
            run.italic = True
        elif seg.startswith("**") and seg.endswith("**"):
            run = para.add_run(seg[2:-2])
            run.bold = True
        elif seg.startswith("*") and seg.endswith("*"):
            run = para.add_run(seg[1:-1])
            run.italic = True
        else:
            para.add_run(seg)