"""
Q1 合成测试 v2：用 python-docx 生成已知内容文档 → 提取 → 对比
docx 对中文支持好，可精确控制表格/段落内容
"""
import json
import os
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ── 测试用例 ──

def make_table_simple(doc: Document):
    """3行3列中文表格"""
    doc.add_heading("第一季度销售数据汇总", level=1)
    doc.add_paragraph("以下为本季度各区域销售业绩概览：")
    table = doc.add_table(rows=4, cols=3, style="Table Grid")
    headers = ["区域", "Q1 销售额（万元）", "同比增长"]
    data = [
        ["华东", "1,245.8", "12.3%"],
        ["华南", "987.2", "7.1%"],
        ["华北", "1,567.0", "15.8%"],
    ]
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            table.cell(i + 1, j).text = val
    doc.add_paragraph("")


def make_table_complex(doc: Document):
    """5列4行含百分比和状态"""
    doc.add_heading("各业务线收入构成", level=1)
    table = doc.add_table(rows=5, cols=5, style="Table Grid")
    headers = ["业务线", "金额（亿）", "占比", "环比变化", "状态"]
    data = [
        ["云计算", "78.6", "42.3%", "+5.2pp", "达标"],
        ["AI 平台", "45.2", "24.3%", "+8.1pp", "超额"],
        ["数据服务", "32.1", "17.3%", "-2.1pp", "预警"],
        ["安全合规", "29.8", "16.1%", "+1.5pp", "达标"],
    ]
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            table.cell(i + 1, j).text = val
    doc.add_paragraph("")


def make_text_headings(doc: Document):
    """多级标题 + 正文段落"""
    doc.add_heading("新能源电池技术发展报告", level=0)

    doc.add_heading("1 锂电池技术现状", level=1)
    doc.add_paragraph(
        "锂电池作为当前主流的动力电池技术，能量密度在过去十年间提升了约 3 倍。"
        "特斯拉 4680 电池采用了无极耳设计，将能量密度提升至 300Wh/kg 以上。"
    )

    doc.add_heading("2 固态电池产业化进展", level=1)
    doc.add_paragraph(
        "固态电池被视为下一代电池技术的关键方向。目前主要面临界面阻抗、"
        "制造成本和规模化生产三大挑战。宁德时代、三星 SDI 等头部企业"
        "均计划在 2027-2028 年实现小批量量产。"
    )

    doc.add_heading("3 钠离子电池的机遇", level=1)
    doc.add_paragraph(
        "钠离子电池在储能和低速电动车领域展现出成本优势。与锂电池相比，"
        "钠资源的地壳丰度是锂的 1000 倍以上，供应链风险显著降低。"
    )


# ── 预期检查项 ──

CHECK_TABLE_SIMPLE = {
    "name": "简单表格 3x3",
    "is_table": True,
    "must_contain": [
        "华东", "华南", "华北",
        "1,245.8", "987.2", "1,567.0",
        "12.3%", "7.1%", "15.8%",
    ],
    "headers_expected": ["区域", "销售额", "同比增长"],
}

CHECK_TABLE_COMPLEX = {
    "name": "复杂表格 5x4",
    "is_table": True,
    "must_contain": [
        "云计算", "AI 平台", "数据服务", "安全合规",
        "78.6", "45.2", "32.1", "29.8",
        "42.3%", "24.3%",
        "达标", "超额", "预警",
    ],
    "headers_expected": ["业务线", "金额", "占比"],
}

CHECK_TEXT_HEADINGS = {
    "name": "多级标题 + 正文",
    "is_table": False,
    "must_contain": [
        "锂电池", "固态电池", "钠离子电池",
        "4680", "300Wh/kg",
        "宁德时代", "三星", "SDI",
        "2027-2028",
        "1000 倍", "界面阻抗",
    ],
    "must_have_sections": [
        "锂电池技术现状", "固态电池产业化进展", "钠离子电池的机遇",
    ],
}


def run_test(doc_builder, checks: dict):
    """生成 DOCX → 提取 → 对比 → 返回结果"""
    with tempfile.TemporaryDirectory() as tmpdir:
        doc = Document()
        # 设置默认字体支持中文
        style = doc.styles["Normal"]
        font = style.font
        font.name = "SimSun"
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

        doc_builder(doc)

        docx_path = os.path.join(tmpdir, "test.docx")
        doc.save(docx_path)
        size_kb = os.path.getsize(docx_path) / 1024

        # 提取
        from agent_file_create.document.extractor import extract_from_file
        result = extract_from_file(str(docx_path))
        if isinstance(result, dict):
            extracted_dict = result
        elif hasattr(result, "dict"):
            extracted_dict = result.dict()
        else:
            extracted_dict = {"_raw": str(result)}
        extracted_str = json.dumps(extracted_dict, ensure_ascii=False, default=str).lower()

        # 对比
        passes: list[str] = []
        fails: list[str] = []

        for text in checks.get("must_contain", []):
            if text.lower() in extracted_str:
                passes.append(f"含 '{text}'")
            else:
                fails.append(f"缺 '{text}'")

        for section in checks.get("must_have_sections", []):
            if section.lower() in extracted_str:
                passes.append(f"章节 '{section}'")
            else:
                fails.append(f"缺章节 '{section}'")

        for h in checks.get("headers_expected", []):
            if h.lower() in extracted_str:
                passes.append(f"表头 '{h}'")
            else:
                fails.append(f"缺表头 '{h}'")

        score = len(passes) / max(len(passes) + len(fails), 1) * 100

        return {
            "name": checks["name"],
            "pass": len(fails) == 0,
            "score_pct": round(score, 1),
            "passes": len(passes),
            "failures": len(fails),
            "details_pass": passes,
            "details_fail": fails,
            "size_kb": size_kb,
            "has_error": "error" in extracted_dict,
            "error_msg": str(extracted_dict.get("error", ""))[:200],
            "extraction_preview": extracted_str[:500],
        }


def main():
    cases = [
        (make_table_simple, CHECK_TABLE_SIMPLE),
        (make_table_complex, CHECK_TABLE_COMPLEX),
        (make_text_headings, CHECK_TEXT_HEADINGS),
    ]

    results = []
    for builder, checks in cases:
        print(f"\n{'='*60}")
        print(f"Test: {checks['name']}")
        r = run_test(builder, checks)
        results.append(r)

        status = "PASS" if r["pass"] else "FAIL"
        if r.get("has_error"):
            status += " (extract error)"
        print(f"  -> {status} | Score {r['score_pct']}% | +{r['passes']} / -{r['failures']}")
        if r.get("error_msg"):
            print(f"  Extract error: {r['error_msg']}")
        for f in r["details_fail"]:
            print(f"    [X] {f}")
        print(f"  Preview: {r['extraction_preview'][:300]}")

    print(f"\n{'='*60}")
    print("Summary")
    print(f"{'='*60}")
    passed = sum(1 for r in results if r["pass"])
    total = len(results)
    avg = sum(r["score_pct"] for r in results) / max(total, 1)
    print(f"Passed: {passed}/{total} | Avg Score: {avg:.1f}%")
    for r in results:
        icon = "[PASS]" if r["pass"] else "[FAIL]"
        print(f"  {icon} {r['name']}: {r['score_pct']}%")

    out = os.path.join(os.path.dirname(__file__), "..", "result", "q1_synthetic_test.json")
    os.makedirs(os.path.dirname(out), exist_ok=True)
    with open(out, "w", encoding="utf-8") as f:
        json.dump({"results": results, "summary": {"passed": passed, "total": total, "avg_score": avg}}, f, ensure_ascii=False, indent=2)
    print(f"\nDetail: {out}")


if __name__ == "__main__":
    main()
