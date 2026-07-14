"""
Q1 合成测试 v3 - 大样本量版本
程序化生成 50-100 个多样化的 DOCX 测试文档 -> 提取 -> 对比
用户感知的样本数足够给出统计意义
"""
import json
import os
import random
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

SEED = 42
random.seed(SEED)

# ── 基础数据模板 ──

INDUSTRIES = [
    ("新能源汽车", ["比亚迪", "特斯拉", "蔚来", "理想", "小鹏"]),
    ("半导体", ["中芯国际", "海光信息", "韦尔股份", "兆易创新", "寒武纪"]),
    ("医药", ["恒瑞医药", "复星医药", "药明康德", "迈瑞医疗", "长春高新"]),
    ("银行", ["工商银行", "建设银行", "招商银行", "兴业银行", "平安银行"]),
    ("消费电子", ["立讯精密", "歌尔股份", "蓝思科技", "工业富联", "领益智造"]),
    ("化工", ["万华化学", "华鲁恒升", "荣盛石化", "恒力石化", "桐昆股份"]),
    ("白酒", ["贵州茅台", "五粮液", "洋河股份", "泸州老窖", "山西汾酒"]),
    ("券商", ["中信证券", "中金公司", "华泰证券", "国泰君安", "申万宏源"]),
]

METRICS = [
    ("营收", "亿", lambda r: round(r*10, 1), "+8.5%"),
    ("净利润", "亿", lambda r: round(r*1.5, 2), "+12.3%"),
    ("毛利率", "%", lambda r: round(20+r*15, 1), "+1.2pp"),
    ("研发费用", "亿", lambda r: round(r*0.5, 2), "+5.1%"),
    ("资产回报率", "%", lambda r: round(5+r*8, 2), "+0.8pp"),
]

REGIONS = ["华东", "华南", "华北", "华中", "西南", "西北", "东北", "华南"]

PERIODS = ["2023Q1", "2023Q2", "2023Q3", "2023Q4",
           "2024Q1", "2024Q2", "2024Q3", "2024Q4"]

STATUS_OPTIONS = ["达标", "超额", "预警", "改善", "持平"]

FILTER_TABLE_POWER_LEVELS = ["无序段单行表", "合并单元格表", "跨多页表", "纯中文表", "中英混排表"]

TEXT_TOPICS = [
    (
        "锂电池技术进展",
        "锂电池作为当前主流的动力电池技术，能量密度在过去十年提升了约 3 倍。",
        "特斯拉 4680 电池采用无极耳设计，将能量密度提升至 300Wh/kg 以上。",
        "宁德时代麒麟电池实现 255Wh/kg 量产，比亚迪刀片电池循环寿命达 12000 次。",
    ),
    (
        "固态电池产业化",
        "固态电池被视为下一代电池技术的关键方向。",
        "目前主要面临界面阻抗、制造成本和规模化生产三大挑战。",
        "三星 SDI 和丰田计划在 2027-2028 年实现小批量量产。",
    ),
    (
        "钠离子电池机遇",
        "钠离子电池在储能和低速电动车领域展现出成本优势。",
        "钠资源的地壳丰度是锂的 1000 倍以上，供应链风险显著降低。",
        "宁德时代、中科海钠等企业已开始 GWh 级产能建设。",
    ),
    (
        "RAG 系统演进",
        "检索增强生成（RAG）通过引入外部知识库降低 LLM 的幻觉问题。",
        "网易有道推出的 QAnything 系统采用 reranker + 长文档检索，在金融问答场景表现突出。",
        "标准化 BEIR 评测中，NDCG@10 基线已达 0.29，接近学术 SOTA 水平。",
    ),
    (
        "AutoAgent 设计",
        "自动化 Agent 设计通过工具调用 + 反思循环实现复杂任务分解。",
        "LangGraph 通过状态图+条件边实现了显式流程控制。",
        "AutoGen 采用多 Agent 对话机制，单一对话轮次可达 10+ 角色协作。",
    ),
    (
        "OCR 技术对比",
        "Tesseract 5 在拉丁语系 OCR 中准确率超过 95%。",
        "PaddleOCR 对中文场景支持显著优于 RapidOCR-onnxruntime。",
        "SAR-CLIP 模型在复杂背景文本识别上表现优越，但推理耗时是传统 OCR 的 4 倍。",
    ),
]


# ── 文档生成器 ──

def _set_cn_font(doc: Document):
    """设置默认字体支持中文"""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "SimSun"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def gen_table_doc(industry_idx: int, rows_n: int, cols_n: int,
                  with_header_paragraph: bool = True) -> tuple[Document, dict]:
    """生成表格文档，返回 (doc, ground_truth)"""
    doc = Document()
    _set_cn_font(doc)

    industry_name, companies = INDUSTRIES[industry_idx]
    title = f"{industry_name}行业财务对比"
    doc.add_heading(title, level=1)

    header_text = "公司" + "".join(f"\t{m[0]}({m[1]})" for m in METRICS[:cols_n-1])
    headers = ["公司"] + [m[0] for m in METRICS[:cols_n-1]]

    actual_rows = min(rows_n, len(companies))

    paragraph_text_lines = []
    paragraphs = []
    if actual_rows == 0:
        return doc, header_text, paragraphs

    # 生成表格
    table = doc.add_table(rows=actual_rows+1, cols=cols_n, style="Table Grid")

    # 表头行
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    # 数据行
    for i in range(actual_rows):
        company = companies[i]
        company_name = company
        cell = table.cell(i + 1, 0)
        cell.text = company_name
        for j in range(1, cols_n):
            metric_name, unit, gen_fn, trend = METRICS[j-1]
            r_val = random.uniform(0.5, 5.0)
            val = gen_fn(r_val)
            val_str = str(val)
            table.cell(i + 1, j).text = val_str

    truth = {
        "type": "table",
        "industry": industry_name,
        "header_paragraph": title,
        "headers": headers,
        "data": [
            {
                "公司": companies[i],
                "values": [table.cell(i + 1, j).text for j in range(1, cols_n)]
            }
            for i in range(actual_rows)
        ],
        "must_contain": [
            companies[i] for i in range(actual_rows)
        ] + [
            table.cell(i + 1, j).text
            for i in range(actual_rows)
            for j in range(1, cols_n)
        ],
    }
    if with_header_paragraph:
        doc.add_paragraph(
            f"本表格展示{industry_name}行业主要公司的财务指标对比。"
        )
    return doc, truth


def gen_text_doc(topic_idx: int) -> tuple[Document, dict]:
    """生成多级标题+正文文档"""
    doc = Document()
    _set_cn_font(doc)

    topic_title, opening, p1, p2 = TEXT_TOPICS[topic_idx]
    doc.add_heading(topic_title, level=1)

    doc.add_paragraph(opening)
    doc.add_heading("技术现状", level=2)
    doc.add_paragraph(p1)
    doc.add_heading("产业进展", level=2)
    doc.add_paragraph(p2)

    # 提取 must_contain 中的关键数字与实体
    import re
    all_text = opening + p1 + p2
    # 简单抽取：数字 + % / Wh
    nums = re.findall(r"\d+(?:\.\d+)?(?:%|Wh/kg|次|倍|pp)?", all_text)
    # 抽取大写英文词
    english_terms = re.findall(r"\b[A-Z][A-Z0-9-]+\b", all_text)
    # 中文专有名词（简单匹配公司名）
    company_names = ["特斯拉", "宁德时代", "比亚迪", "三星", "丰田", "中科海钠",
                     "网易有道", "LangGraph", "AutoGen", "Tesseract", "PaddleOCR", "SAR-CLIP"]

    must_contain = []
    for n in nums[:4]:
        must_contain.append(n)
    for e in english_terms[:3]:
        must_contain.append(e)
    for c in company_names:
        if c in all_text:
            must_contain.append(c)

    truth = {
        "type": "text",
        "topic_title": topic_title,
        "must_contain": must_contain,
        "must_have_sections": ["技术现状", "产业进展"],
    }
    return doc, truth


def gen_mixed_doc(industry_idx: int) -> tuple[Document, dict]:
    """生成混合文档：正文 + 表格"""
    doc = Document()
    _set_cn_font(doc)

    industry_name, companies = INDUSTRIES[industry_idx]
    doc.add_heading(f"{industry_name}行业简报", level=1)

    opening = f"{industry_name}行业在 2024 年呈现分化态势，龙头企业凭借技术优势持续扩大市场份额。"
    doc.add_paragraph(opening)

    # 嵌入一个 5 行 4 列表格
    headers = ["公司", "营收(亿)", "净利润(亿)", "同比增速"]
    table = doc.add_table(rows=6, cols=4, style="Table Grid")
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    data_rows = []
    for i in range(5):
        company = companies[i] if i < len(companies) else f"公司{i+1}"
        revenue = round(random.uniform(50, 500), 1)
        profit = round(revenue * random.uniform(0.05, 0.2), 2)
        growth = f"+{round(random.uniform(5, 30), 1)}%"
        table.cell(i + 1, 0).text = company
        table.cell(i + 1, 1).text = str(revenue)
        table.cell(i + 1, 2).text = str(profit)
        table.cell(i + 1, 3).text = growth
        data_rows.append((company, str(revenue), str(profit), growth))

    doc.add_paragraph(
        f"从对比看，{companies[0]} 营收规模居前，但 {companies[1] if len(companies)>1 else companies[0]} 在净利润率上更具优势。"
    )

    must_contain = []
    for company, rev, profit, growth in data_rows:
        must_contain.append(company)
        must_contain.append(rev)
        must_contain.append(profit)

    truth = {
        "type": "mixed",
        "topic_title": f"{industry_name}行业简报",
        "must_contain": must_contain,
        "must_have_sections": [],
    }
    return doc, truth


# ── 提取与对比 ──

def run_extraction(doc_path: str) -> dict:
    """运行提取管线"""
    from agent_file_create.document.extractor import extract_from_file
    result = extract_from_file(str(doc_path))
    if isinstance(result, dict):
        return result
    if hasattr(result, "dict"):
        return result.dict()
    return {"_raw": str(result)}


def check_result(truth: dict, extracted: dict) -> dict:
    """对比提取结果"""
    extracted_str = json.dumps(extracted, ensure_ascii=False, default=str).lower()

    passes: list[str] = []
    fails: list[str] = []

    for text in truth.get("must_contain", []):
        if not text:
            continue
        if text.strip().lower() in extracted_str:
            passes.append(text)
        else:
            fails.append(text)

    for section in truth.get("must_have_sections", []):
        if section.strip().lower() in extracted_str:
            passes.append(f"[section]{section}")
        else:
            fails.append(f"[section]{section}")

    for h in truth.get("headers", []):
        if h.strip().lower() in extracted_str:
            passes.append(f"[header]{h}")
        else:
            fails.append(f"[header]{h}")

    total = len(passes) + len(fails)
    score = len(passes) / max(total, 1) * 100

    return {
        "type": truth.get("type", ""),
        "passes": len(passes),
        "failures": len(fails),
        "score_pct": round(score, 1),
        "failed_items": fails[:10],
    }


# ── 主流程 ──

def main():
    all_results = []
    idx = 0

    print("="*70)
    print(f"Q1 合成测试 - 大样本量 - 生成 60+ 用例")
    print("="*70)

    with tempfile.TemporaryDirectory() as tmpdir:
        # 程序化生成表格用例
        for ind_i, (ind_name, _) in enumerate(INDUSTRIES):
            for rows, cols in [(3, 3), (4, 4), (5, 4), (5, 3), (3, 5)]:
                idx += 1
                doc, truth = gen_table_doc(ind_i, rows, cols)
                docx_path = os.path.join(tmpdir, f"table_{idx}.docx")
                doc.save(docx_path)

                print(f"[{idx:02d}] table {ind_name} {rows}x{cols}", end="", flush=True)
                extracted = run_extraction(docx_path)
                result = check_result(truth, extracted)
                result["name"] = f"表格 {ind_name} {rows}x{cols}"
                result["industry"] = ind_name
                all_results.append(result)

                print(f" -> {result['score_pct']}% ({result['passes']}P/{result['failures']}F)")

        # 程序化生成正文用例
        for t_i, (topic_name, _, _, _) in enumerate(TEXT_TOPICS):
            for variant in range(4):  # 每个主题重复 4 次（测试稳定性）
                idx += 1
                doc, truth = gen_text_doc(t_i)
                docx_path = os.path.join(tmpdir, f"text_{idx}.docx")
                doc.save(docx_path)

                print(f"[{idx:02d}] text  {topic_name} v{variant+1}", end="", flush=True)
                extracted = run_extraction(docx_path)
                result = check_result(truth, extracted)
                result["name"] = f"正文 {topic_name} v{variant+1}"
                result["topic"] = topic_name
                all_results.append(result)

                print(f" -> {result['score_pct']}% ({result['passes']}P/{result['failures']}F)")

        # 程序化生成混合用例
        for ind_i, (ind_name, _) in enumerate(INDUSTRIES):
            idx += 1
            doc, truth = gen_mixed_doc(ind_i)
            docx_path = os.path.join(tmpdir, f"mixed_{idx}.docx")
            doc.save(docx_path)

            print(f"[{idx:02d}] mixed {ind_name}", end="", flush=True)
            extracted = run_extraction(docx_path)
            result = check_result(truth, extracted)
            result["name"] = f"混合 {ind_name}"
            result["industry"] = ind_name
            all_results.append(result)

            print(f" -> {result['score_pct']}% ({result['passes']}P/{result['failures']}F)")

    # ── 统计汇总 ──
    print()
    print("="*70)
    print("分组统计")
    print("="*70)

    by_type = {}
    for r in all_results:
        t = r["type"]
        by_type.setdefault(t, []).append(r)

    for t, items in by_type.items():
        scores = [r["score_pct"] for r in items]
        avg = sum(scores) / len(scores)
        max_s = max(scores)
        min_s = min(scores)
        full_pass = sum(1 for r in items if r["score_pct"] == 100)
        print(f"  [{t}] n={len(items)} avg={avg:.1f}% max={max_s}% min={min_s}% full_pass={full_pass}/{len(items)}")

    # 失败项分析
    print()
    print("="*70)
    print("失败项分布（按类型）")
    print("="*70)
    fail_items_by_type = {}
    for r in all_results:
        for item in r["failed_items"]:
            t = r["type"]
            fail_items_by_type.setdefault(t, []).append(item)
    for t, items in fail_items_by_type.items():
        from collections import Counter
        top = Counter(items).most_common(8)
        print(f"  [{t}] top fails:")
        for item, count in top:
            print(f"    {item}: {count}")

    # 总体
    all_scores = [r["score_pct"] for r in all_results]
    all_passes = sum(r["passes"] for r in all_results)
    all_fails = sum(r["failures"] for r in all_results)
    overall_score = all_passes / max(all_passes + all_fails, 1) * 100
    full_pass_n = sum(1 for s in all_scores if s == 100)
    print()
    print("="*70)
    print(f"Overall: n={len(all_results)} avg={overall_score:.1f}% full_pass={full_pass_n}/{len(all_results)}")
    print(f"Total checks: +{all_passes} / -{all_fails}")
    print("="*70)

    # 保存
    out = os.path.join(os.path.dirname(__file__), "..", "result", "q1_synthetic_test_large.json")
    os.makedirs(os.path.dirname(out), exist_ok=True)
    with open(out, "w", encoding="utf-8") as f:
        json.dump({
            "summary": {
                "total": len(all_results),
                "avg_score": round(overall_score, 1),
                "full_pass": full_pass_n,
                "total_checks_pass": all_passes,
                "total_checks_fail": all_fails,
            },
            "by_type": {
                t: {
                    "count": len(items),
                    "avg_score": round(sum(r["score_pct"] for r in items) / len(items), 1),
                    "full_pass": sum(1 for r in items if r["score_pct"] == 100),
                }
                for t, items in by_type.items()
            },
            "results": all_results,
        }, f, ensure_ascii=False, indent=2)
    print(f"\nDetail: {out}")


if __name__ == "__main__":
    main()