"""Generate the project application .docx file."""
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── Helper functions ──
def add_paragraph(text, bold=False, size=12, font_name='宋体', alignment=None, indent_cm=None, spacing_after=6, first_line_indent_cm=None):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(spacing_after)
    pf.space_before = Pt(0)
    if indent_cm is not None:
        pf.left_indent = Cm(indent_cm)
    if first_line_indent_cm is not None:
        pf.first_line_indent = Cm(first_line_indent_cm)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold
    return p

def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def merge_cells(table, row_start, col_start, row_end, col_end):
    cell = table.cell(row_start, col_start)
    cell_end = table.cell(row_end, col_end)
    cell.merge(cell_end)
    return cell

def add_cell_text(cell, text, bold=False, size=12, font_name='宋体', alignment=None, v_alignment=None):
    """Set text in a cell, clearing existing paragraphs."""
    for pp in cell.paragraphs:
        for run in pp.runs:
            run.text = ''
    p = cell.paragraphs[0]
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(18)
    if v_alignment is not None:
        cell.vertical_alignment = v_alignment
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.bold = bold

def add_cell_paragraphs(cell, texts, bold=False, size=12, font_name='宋体', alignment=None):
    """Set multiple text lines in a cell."""
    for i, text in enumerate(texts):
        if i == 0:
            p = cell.paragraphs[0]
            for run in p.runs:
                run.text = ''
        else:
            p = cell.add_paragraph()
        if alignment is not None:
            p.alignment = alignment
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.bold = bold

# ══════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════
add_paragraph('', size=12)
p = add_paragraph('大学生创新训练项目申请书', bold=True, size=18, font_name='黑体',
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=12)
add_paragraph('', size=12, spacing_after=0)

# ══════════════════════════════════════════════════════════════════
# COVER TABLE
# ══════════════════════════════════════════════════════════════════
cover_data = [
    ('项目编号', '', '项目名称', '基于大语言模型的多源信息智能文档生成系统'),
    ('项目负责人', '叶永青', '学  号', '2024463030216'),
    ('所在学院', '计算机科学与技术学院\n(软件学院、网络空间安全学院)', '专业班级', '计算机类'),
    ('联系电话', '', '指导教师', '陶铭, 刘群'),
    ('E-mail', 'taoming6723@126.com, 472418964@qq.com', '', ''),
    ('申请日期', '2025年04月27日', '起止年月', '2025年04月至2027年04月'),
]

table1 = doc.add_table(rows=len(cover_data), cols=4)
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
table1.style = 'Table Grid'
# Set column widths
for row_idx, (label1, val1, label2, val2) in enumerate(cover_data):
    row = table1.rows[row_idx]
    for cell in row.cells:
        cell._tc.get_or_add_tcPr()
    # Set widths
    row.cells[0].width = Cm(2.8)
    row.cells[1].width = Cm(4.0)
    row.cells[2].width = Cm(2.0)
    row.cells[3].width = Cm(4.2)

    add_cell_text(row.cells[0], label1, bold=False, size=12, font_name='黑体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(row.cells[1], val1, size=12, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(row.cells[2], label2, bold=False, size=12, font_name='黑体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(row.cells[3], val2, size=12, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Set row height
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="600" w:hRule="atLeast"/>')
    trPr.append(trHeight)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 填写说明
# ══════════════════════════════════════════════════════════════════
add_paragraph('填 写 说 明', bold=True, size=18, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=12)
add_paragraph('1、申请书所列各项内容均须实事求是填写，表达明确严谨，简明扼要。模板可网上下载、自行加页。',
              size=12, font_name='仿宋_GB2312', first_line_indent_cm=0.74, spacing_after=3)
add_paragraph('2、申请书首页只填写项目负责人。"项目编号"一栏可不填。',
              size=12, font_name='仿宋_GB2312', first_line_indent_cm=0.74, spacing_after=3)
add_paragraph('3、项目负责人所在院系须认真审核，签署推荐意见并加盖公章后提交。',
              size=12, font_name='仿宋_GB2312', first_line_indent_cm=0.74, spacing_after=12)

# ══════════════════════════════════════════════════════════════════
# 一、基本情况
# ══════════════════════════════════════════════════════════════════
add_paragraph('一、基本情况', bold=True, size=14, font_name='黑体', spacing_after=6)

basic_table = doc.add_table(rows=9, cols=10)
basic_table.alignment = WD_TABLE_ALIGNMENT.CENTER
basic_table.style = 'Table Grid'

# Configure columns
col_widths = [Cm(1.1), Cm(1.1), Cm(2.0), Cm(1.4), Cm(1.0), Cm(1.2), Cm(1.2), Cm(1.6), Cm(1.2), Cm(1.7)]

# Row 0: 项目名称
merge_cells(basic_table, 0, 0, 0, 1)
merge_cells(basic_table, 0, 2, 0, 9)
add_cell_paragraphs(basic_table.cell(0, 0), ['项目', '名称'], size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(basic_table.cell(0, 2), '基于大语言模型的多源信息智能文档生成系统', size=10, font_name='楷体_GB2312', alignment=WD_ALIGN_PARAGRAPH.CENTER)

# Row 1: 所属学科
merge_cells(basic_table, 1, 0, 1, 1)
merge_cells(basic_table, 1, 2, 1, 4)
merge_cells(basic_table, 1, 5, 1, 7)
merge_cells(basic_table, 1, 8, 1, 9)
add_cell_paragraphs(basic_table.cell(1, 0), ['所属', '学科'], size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(basic_table.cell(1, 2), '学科一级门：工学', size=10, font_name='宋体')
add_cell_text(basic_table.cell(1, 5), '学科二级类：计算机类', size=10, font_name='宋体')
add_cell_text(basic_table.cell(1, 8), '', size=10, font_name='宋体')

# Row 2: 申请金额
merge_cells(basic_table, 2, 0, 2, 1)
merge_cells(basic_table, 2, 2, 2, 4)
merge_cells(basic_table, 2, 5, 2, 6)
merge_cells(basic_table, 2, 7, 2, 9)
add_cell_paragraphs(basic_table.cell(2, 0), ['申请', '金额'], size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(basic_table.cell(2, 2), '20000  元', size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(basic_table.cell(2, 5), '起止年月', size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(basic_table.cell(2, 7), '2025年04月至2027年04月', size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)

# Row 3: 负责人姓名
merge_cells(basic_table, 3, 0, 3, 1)
merge_cells(basic_table, 3, 2, 3, 3)
merge_cells(basic_table, 3, 4, 3, 5)
merge_cells(basic_table, 3, 7, 3, 9)
add_cell_paragraphs(basic_table.cell(3, 0), ['负责人', '姓名'], size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(basic_table.cell(3, 2), '叶永青', size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(basic_table.cell(3, 4), '性别', size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(basic_table.cell(3, 6), '民族', size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(basic_table.cell(3, 7), '出生年月', size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)

# Row 4: 学号
merge_cells(basic_table, 4, 0, 4, 1)
merge_cells(basic_table, 4, 2, 4, 3)
merge_cells(basic_table, 4, 4, 4, 5)
merge_cells(basic_table, 4, 6, 4, 9)
add_cell_paragraphs(basic_table.cell(4, 0), ['学号'], size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(basic_table.cell(4, 2), '2024463030216', size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_paragraphs(basic_table.cell(4, 4), ['联系', '电话'], size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(basic_table.cell(4, 6), '手机：', size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.LEFT)

# Row 5: 指导教师 row 1 (陶铭)
merge_cells(basic_table, 5, 0, 5, 1)
merge_cells(basic_table, 5, 2, 5, 3)
merge_cells(basic_table, 5, 4, 5, 5)
merge_cells(basic_table, 5, 6, 5, 9)
add_cell_paragraphs(basic_table.cell(5, 0), ['指导', '教师'], size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(basic_table.cell(5, 2), '陶铭', size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_paragraphs(basic_table.cell(5, 4), ['联系', '电话'], size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(basic_table.cell(5, 6), '手机：13763260593', size=10, font_name='宋体')

# Row 6: 指导教师 row 2 (刘群)
merge_cells(basic_table, 6, 0, 6, 1)
merge_cells(basic_table, 6, 2, 6, 3)
merge_cells(basic_table, 6, 4, 6, 5)
merge_cells(basic_table, 6, 6, 6, 9)
add_cell_paragraphs(basic_table.cell(6, 0), ['指导', '教师'], size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(basic_table.cell(6, 2), '刘群', size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_paragraphs(basic_table.cell(6, 4), ['联系', '电话'], size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_text(basic_table.cell(6, 6), '手机：13926829506', size=10, font_name='宋体')

# Row 7: 负责人曾经参与科研的情况
merge_cells(basic_table, 7, 0, 7, 1)
merge_cells(basic_table, 7, 2, 7, 9)
add_cell_paragraphs(basic_table.cell(7, 0), ['负责人', '曾经参与', '科研的', '情况'], size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_paragraphs(basic_table.cell(7, 2), [
    '项目负责人现为实验室成员，参与实验室日常研究活动。',
    '曾系统学习Python、JavaScript等编程语言，具有Web全栈开发经验；',
    '熟悉大语言模型API调用与Prompt Engineering技术；',
    '已完成agent-file-create智能文档生成系统的技术原型开发（约8000+行Python代码，涵盖LangGraph智能体、RAG知识库、长文生成等模块）。'
], size=10, font_name='宋体')

# Row 8: 指导教师承担科研课题情况
merge_cells(basic_table, 8, 0, 8, 1)
merge_cells(basic_table, 8, 2, 8, 9)
add_cell_paragraphs(basic_table.cell(8, 0), ['指导教', '师承担', '科研课', '题情况'], size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_cell_paragraphs(basic_table.cell(8, 2), [
    '陶铭，博士，教授，东莞市无线传感网络系统重点实验室主任，广东工业大学校外硕士研究生导师，中国计算机学会（CCF）高级会员。主要从事异构网络融合、物联网、云计算关键技术研究及应用系统开发。目前已主持国家自然科学基金（青年基金）、广东省科技计划、广东省自然科学基金等多项科研项目，发表学术论文60余篇，其中第一作者/通讯作者SCI索引论文18篇，申请发明专利10件（已授权5件），获得软件著作权登记10余项。担任国家自然科学基金通信评审专家、广东省科技厅入库专家。',
], size=10, font_name='宋体')

# Set row heights for rows 7 and 8
for row_idx in range(len(basic_table.rows)):
    tr = basic_table.rows[row_idx]._tr
    trPr = tr.get_or_add_trPr()
    if row_idx >= 7:
        trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="1200" w:hRule="atLeast"/>')
    else:
        trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="500" w:hRule="atLeast"/>')
    trPr.append(trHeight)

add_paragraph('')

# ══════════════════════════════════════════════════════════════════
# 二、项目主要人员
# ══════════════════════════════════════════════════════════════════
add_paragraph('二、项目主要人员', bold=True, size=14, font_name='黑体', spacing_after=6)

member_table = doc.add_table(rows=4, cols=6)
member_table.alignment = WD_TABLE_ALIGNMENT.CENTER
member_table.style = 'Table Grid'

# Header row
headers = ['姓名', '学号', '专业班级', '所在学院', '项目中的分工', '签名']
for i, h in enumerate(headers):
    add_cell_text(member_table.cell(0, i), h, size=10, font_name='宋体', bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

# Team members
members = [
    ('叶永青', '2024463030216', '计算机类', '计算机科学与技术学院\n(软件学院、网络空间安全学院)', '项目负责人\n系统架构设计与核心开发', ''),
    ('（待补充）', '', '', '', '', ''),
    ('（待补充）', '', '', '', '', ''),
]
for row_idx, member in enumerate(members):
    for col_idx, val in enumerate(member):
        add_cell_text(member_table.cell(row_idx + 1, col_idx), val, size=10, font_name='宋体', alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_paragraph('')

# ══════════════════════════════════════════════════════════════════
# 二（续）、立项依据
# ══════════════════════════════════════════════════════════════════
add_paragraph('三、立项依据（可加页）', bold=True, size=14, font_name='黑体', spacing_after=6)

# Content sections - use the project proposal content, condensed for the application form
sections_text = {
    '（一）项目简介': (
        '本项目旨在设计并实现一个基于大语言模型（LLM）的多源信息智能文档生成系统。'
        '系统接收用户上传的PDF、Word、PPT、Excel、图片等多种格式材料，自动完成信息抽取、大纲规划、'
        '分章节正文生成与模板渲染，输出结构化的Markdown/Word/PDF报告。核心研究内容包括：'
        '（1）多模态文件信息的并行抽取与语义融合；'
        '（2）基于LangGraph的ReAct智能体任务编排；'
        '（3）融合HyDE假设文档嵌入与思维链推理的RAG知识增强检索；'
        '（4）长文生成中的幻觉抑制与语义逻辑链保持机制。'
        '系统采用FastAPI + LangChain技术栈，具备流式对话交互、实时生成预览和知识库管理能力。'
    ),
    '（二）研究目的': (
        '在信息化高速发展的今天，企业和个人每天产生海量文档数据——市场调研、竞品分析、会议纪要、'
        '技术方案、学术综述等。撰写高质量报告通常需要经历"资料收集→信息提取→大纲构思→逐章撰写→审核修改"'
        '的冗长过程，耗时数小时甚至数天。大语言模型（LLM）的出现为自动化文本生成带来了革命性可能，但直接'
        '将LLM应用于专业文档生成仍面临四大核心挑战：一是多源异构信息难以融合，原始材料格式多样（PDF扫描件、'
        'Word文档、Excel数据表、PPT演示、图片截图），传统方法难以统一抽取和关联；二是检索增强生成（RAG）'
        '对推理类问题效果差，用户查询与知识库文档间存在"词汇鸿沟"；三是长文生成存在幻觉与逻辑断裂，LLM在'
        '生成数千字多章节报告时容易编造具体数字和机构名，且章节间缺乏因果递进关系；四是人机协作机制不足，'
        '现有LLM应用多为"一次性输入-输出"模式，缺乏对文档生成全流程的精细控制。'
        '针对上述问题，本项目的研究目标是设计并实现一个面向多源信息输入、具备推理能力和知识增强的智能文档'
        '生成系统，实现从材料上传到正式报告的端到端自动化，并在幻觉抑制、逻辑连贯性、人机交互体验上达到实用水平。'
    ),
    '（三）研究内容': (
        '（1）多格式文件统一信息抽取。收集并整理多种常见办公文档格式（PDF扫描件与电子文档、Word文档、PPT演示文稿、'
        'Excel数据表格、PNG/JPG图片截图等）的测试样本，确保样本覆盖不同排版风格、中文/英文混排、表格嵌套等典型复杂场景。'
        '针对PDF文档，采用PyMuPDF库对电子PDF的文字层进行结构化提取，同时利用其内置的版面分析能力识别段落、标题、'
        '表格区域的边界框坐标；对扫描件及图片型PDF，引入RapidOCR引擎进行中英文混合文字识别，并调用视觉语言模型'
        '（如minicpm-v:8b）对包含复杂图表、流程图、统计图的页面进行语义级别的看图理解，将图表内容转化为自然语言描述。'
        '针对Word文档，基于python-docx解析段落样式层级（标题/正文/列表），提取表格数据并保留行列结构关系。'
        '针对PPT演示文稿，基于python-pptx按幻灯片页码顺序提取每页的标题、正文文本框、图片及备注内容，'
        '将分散于多页的信息按主题聚类重组。针对Excel数据表格，采用openpyxl读取单元格原始数据，'
        '使用pandas进行描述性统计分析（均值、趋势、极值等），再通过LLM将统计结果转化为自然语言趋势解读。'
        '针对图片，对包含文字信息的截图走OCR路径提取文本，对包含复杂视觉内容的图片走Vision LLM路径进行语义理解。'
        '上述所有格式的抽取器通过ThreadPoolExecutor实现并行调度，最大并行度设为4个文件/批次，'
        '各抽取器返回统一的数据结构（包含title标题、keywords关键词列表、summary摘要、key_points关键论点列表、'
        'data结构化数据、conclusion结论六个字段），便于后续大纲生成模块以一致的接口读取。'
        '同时建立针对各格式提取器的单元测试用例，验证在正常文件、空白文件、损坏文件等边界条件下的行为正确性。'
        '\n（2）基于LangGraph的ReAct智能体任务编排。将文档生成的完整工作流程建模为智能体的工具调用决策链，'
        '通过LangGraph框架提供的状态图机制实现"思考—行动—观察—反思"的ReAct推理循环。首先定义七个面向文档生成'
        '场景的专用工具函数：extract_files负责调度上述多格式文件抽取引擎，将原始文件批量转化为结构化信息；'
        'assess_material基于抽取结果的字段填充率（每个文件满分7个有效字段）评估材料完整度，按filled≥5为"充足"、'
        '3≤filled<5为"一般"、filled<3或解析失败过半为"不足"三档输出质量判断及建议；'
        'generate_outline根据抽取结果和用户需求生成包含#号标题、##主章节（≥3个）和###子节的完整报告大纲，'
        '内置格式校验逻辑（检查标题层级合法性、同级标题数量合理性），校验不通过时自动重试最多三次；'
        'generate_content按照大纲分章节生成报告正文，其中顶层H2章节串行生成以保证行文主线逻辑的连贯性，'
        '底层H3子节通过ThreadPoolExecutor并行生成以提高效率，每个H2章节完成后立即将内容写入中间结果文件供前端拉取；'
        'render_templates将生成的大纲和正文填充到预定义的Word/PDF模板文件（基于docxtpl和python-docx）中，'
        '输出最终文档到result/<task_id>/目录；ask_user在材料评估完成后暂停流程，根据质量等级动态生成2-5个澄清问题'
        '（质量好时询问使用场景、目标受众、篇幅偏好等，质量差时询问是否补充材料、对缺失内容的接受度等），'
        '等待用户回复后继续推进；finish在模板渲染成功后标记任务完成，无需输入参数。'
        '智能体在每一步调用后自动解析工具返回结果，判断下一步动作——若材料不足则触发ask_user，'
        '若抽取完成则推进到assess_material→generate_outline，以此类推，形成自适应决策链。'
        '同时利用LangGraph内置的MemorySaver检查点机制，在每次工具调用前自动将当前状态序列化保存，'
        '当任务因异常中断或用户主动暂停时，可通过thread_id恢复至最近的检查点继续执行。'
        '\n（3）推理性RAG知识增强检索。在系统的对话问答与报告生成场景中，知识库检索的质量直接影响生成内容的'
        '事实准确性和论据充分性。针对标准RAG难以处理推理型查询的问题，本项目分三个层次提升检索效果。'
        '第一层为查询改写与扩展：对短于10字的事实型查询进行关键词扩展（从问题中提取核心实体，追加同义词和相关术语），'
        '对包含"为什么""对比""分析""评估"等关键词的推理型查询启用HyDE假设文档嵌入——让LLM先根据问题生成一段'
        '使用专业知识库语言风格撰写的假设回答（约200-600字），再以该假设回答的embedding向量代替原始问题进行检索，'
        '从而桥接用户日常口语与知识库专业文档之间的词汇使用鸿沟。第二层为三路混合检索与融合排序：'
        '密集向量检索（基于bge-m3嵌入模型进行语义相似度匹配）、词汇检索（基于Elasticsearch或内置BM25算法进行'
        '关键词精确命中）、统计频率检索（基于BM25的TF-IDF变体）各自召回top-20候选，共60条候选通过RRF'
        '（Reciprocal Rank Fusion）算法融合排序，再经Cross-encoder重排序模型（BAAI/bge-reranker-v2-m3）精排取top-8。'
        '第三层为推理增强回答生成：设计五步思维链Prompt模板（问题理解→证据梳理标注来源编号→推理链条说明因果关系→'
        '最终回答给出3-6条要点→自我检查逐条核查每个论断是否有检索片段支撑），'
        '对无支撑的推断强制标注"（推测）"或"（材料未覆盖）"。对于多角度对比类复杂问题（如"从成本、性能和可维护性三个维度对比A和B"），'
        '自动将问题分解为2-4个原子性子问题，每个子问题独立完成检索+回答流程，最后通过综合Prompt融合各子问题的分析结果，'
        '标注不同证据之间的关联（因果/对比/互补），当多个子问题的证据存在矛盾时明确指出并给出最可能的结论。'
        '\n（4）长文生成中的幻觉抑制与语义逻辑链保持。LLM在生成长达数千字的多章节报告时存在两类典型问题：'
        '事实性幻觉（编造具体的百分比数字、机构名称、人名、年份等）和逻辑性断裂（章节之间缺乏因果/递进/转折关系，'
        '读起来像独立短文的机械拼接）。针对这两类问题，本项目构建四层递进式幻觉防御体系：'
        '第1层（生成时约束）——在Prompt中内嵌禁止编造规则，要求模型拒绝复读原材料长句、不得编造材料中未出现的具体数字/机构名/人名、'
        '所有事实性论断必须标注"（据某某材料）"来源；第2层（推导时规范）——对表格数据的处理要求表述为趋势和对比关系而非逐行罗列原文数字、'
        '数据缺失时必须明确说明"材料中暂缺相关数据"而非凭空编造填补空白；第3层（完成后审查）——生成全部章节后，'
        '由LLM以终审角色通读全文，执行跨章节矛盾检测（如同一指标在不同章节出现不一致数值）、无依据断言标记'
        '（逐句核查是否可追溯到原始材料）、术语一致性校验（同一概念在全文中是否使用统一名称）；'
        '第4层（核查时交叉比对）——使用正则表达式从原始材料中提取所有数字、百分比、机构名称、年份的集合，'
        '与生成报告中的对应实体逐项交叉比对，标记无法在材料中找到匹配项的论断为"待人工核实"。'
        '在逻辑链保持方面，摒弃常见的固定字数硬截断做法，改为要求LLM在每章生成后将前序章节内容压缩为包含'
        '"核心论点+关键实体+逻辑角色（铺垫/论证/对比/总结）"三元组的语义摘要，后续章节在生成时将该摘要作为'
        '上下文前缀，使其不仅知道"前面写了什么内容"，更明确"前面的论证链条推进到了哪个阶段"。'
        '此外，生成策略上采用串行H2+并行H3的组合方案：顶层主章节按顺序逐一生成，确保全文论证主线的递进逻辑不被打断；'
        '底层子节在主章节方向确定后并行展开，通过ThreadPoolExecutor提高整体生成吞吐量。'
        '\n（5）人机协作交互机制。传统的LLM应用多为"一次性输入-输出"的黑箱模式，用户无法在生成过程中介入调整，'
        '一旦输出不理想只能清空重新开始。本项目设计全流程的人机协作机制，赋予用户对文档生成的精细控制能力。'
        '在任务启动阶段，智能体完成材料评估后强制触发澄清环节，根据材料质量动态生成2-5个针对性问题'
        '（如"报告使用场景是内部决策、对外汇报还是学术发表？""篇幅偏好3000字精简版还是8000字详尽版？"），'
        '用户可在前端界面逐条回答或选择跳过，系统将用户的澄清信息持久化并注入后续生成流程。'
        '在生成进行阶段，前端采用定时轮询（每2秒请求任务状态接口），一旦检测到进入正文生成阶段，'
        '自动拉取已完成写入的章节结果文件，在报告预览区实时渲染Markdown格式的已完成内容，'
        '未完成章节显示"（生成中…）"的占位标记，让用户全程感知生成进展。'
        '在生成完成后，用户可通过对话面板发送自然语言修改指令（如"请把第三章关于成本的论述改得更详细一些"），'
        '系统内置的意图识别模块基于约40个中文修改关键词（如"修改""调整""删除""补充""重写""缩短""详细""简洁"等）'
        '自动判断用户是否希望修改报告，若匹配成功则提取修改目标（章节名称、修改类型、修改方向），'
        '通过/regen命令执行单章节定向重生成：模糊匹配章节标题定位目标章节，仅重新生成该章节及其子节，'
        '保持其他章节内容不变。此外，全局对话采用Server-Sent Events（SSE）协议实现回复内容的流式推送，'
        '用户无需等待完整回答即可逐token看到生成内容，消除长回复的等待焦虑。'
    ),
    '（四）国、内外研究现状和发展动态': (
        '国外研究现状：LangChain已成为LLM应用开发的事实标准框架，LangGraph在此基础上支持状态图驱动的智能体构建。'
        'Meta的Lewis等人（2020）首次提出RAG框架，将检索系统与大语言模型结合；Gao等人（2023）提出HyDE方法，'
        '通过生成假设文档桥接语义鸿沟；微软的GraphRAG（2024）引入知识图谱增强多跳推理能力。'
        'OpenAI的GPT-4和Anthropic的Claude系列将上下文窗口扩展到128K-200K tokens，使长文档处理成为可能。'
        'Wei等人（2022）提出的Chain-of-Thought方法论被广泛应用于多步推理任务。AutoGPT、MetaGPT、CrewAI等'
        '项目探索了LLM驱动的自主智能体架构。'
        '\n国内研究现状：深度求索（DeepSeek）、智谱AI（ChatGLM）、阿里巴巴（通义千问）等发布了一系列国产大模型，'
        '均提供OpenAI API兼容接口。百度PaddleOCR和RapidAI团队的RapidOCR为中文场景提供了高效的OCR解决方案。'
        '国内有Langchain-Chatchat、FastGPT、Dify等开源RAG项目，但多数聚焦于"问答"场景，缺乏面向"长文生成"的垂直深度优化。'
        '\n现有工作的不足与本项目的切入点：当前主流LLM文档工具在输入格式上多仅支持纯文本或单一格式，本项目设计了五类常见格式的'
        '统一抽取管线；在生成模式上多采用一次性问答模式，本项目通过LangGraph智能体实现七个工具协同调用的完整管线；'
        '在长文质量上缺乏系统性的幻觉控制，本项目构建四层递进式防御体系；在推理能力上标准RAG难以应对因果分析和方案对比，'
        '本项目通过HyDE+思维链+问题分解三层机制实现推理增强；在人机协作上本项目提供了澄清交互、实时预览和定向重生成等精细控制手段。'
    ),
    '（五）创新点与项目特色': (
        '1. 多格式文件统一信息抽取与语义融合。设计统一的六字段提取结果结构，将PDF扫描件、Office文档、Excel数据表、'
        '图片截图等异构输入的抽取结果归一化，为后续大纲生成提供一致的信息底座。并行抽取架构显著提升多文件处理吞吐量。'
        '\n2. HyDE+思维链+问题分解的三阶推理RAG。区别于传统RAG的"查询→检索→回答"单一路径：一阶（查询层）HyDE假设文档嵌入'
        '解决词汇不匹配问题；二阶（推理层）五步思维链Prompt强制显式证据溯源；三阶（综合层）复杂问题自动分解→分别检索→'
        '矛盾检测→综合，覆盖多角度分析场景。'
        '\n3. 四层递进式幻觉防御体系。将幻觉抑制从"事后审查"提升为"全程嵌入"——Prompt内嵌约束（生成时）→结构化指令（推导时）'
        '→LLM终审（完成后）→正则事实交叉比对（核查时），实现对具体数字、机构名、人名等编造行为的系统性拦截。'
        '\n4. LLM语义摘要驱动的章节逻辑衔接。放弃常见的固定字数硬截断做法，独创性地要求LLM将前序章节摘要为包含"核心论点+关键实体+'
        '逻辑角色"的语义描述，使后续章节不仅知道"前面写了什么"，更知道"前面在论证链中处于什么位置"。'
        '\n5. 全流程人机协作机制。打破LLM应用"一次输入-输出"的黑箱模式——智能体在信息不足时主动暂停并生成澄清问题；'
        '生成过程中实时增量预览已完成章节；支持40个中文关键词自动识别用户修改意图和单章节定向重生成。'
    ),
    '（六）技术路线、拟解决的问题及预期成果': (
        '1. 技术路线。本项目的技术方案分为四个层次递进建设。'
        '第一层为基础能力层：构建多格式文件抽取引擎（PDF/PyMuPDF+RapidOCR、Word/python-docx、PPT/python-pptx、'
        'Excel/openpyxl+pandas、图片/OCR+视觉模型），LLM统一工厂同时封装ChatOpenAI和ChatOllama双后端，'
        '向量存储支持SQLite本地开发和PostgreSQL+pgvector生产部署。'
        '第二层为智能体编排层：基于LangGraph构建ReAct智能体，装配七个工具协同调用，形成"思考—行动—观察—反思"闭环，'
        'MemorySaver检查点确保断点可恢复。'
        '第三层为知识增强与推理层：RAG管线实现"查询改写→三路混合召回→RRF融合→Cross-encoder重排序→思维链生成"，'
        '长文生成增强构建四层幻觉防御和LLM语义摘要衔接。'
        '第四层为交互服务层：FastAPI构建全部REST路由并附加SSE流式响应，前端零依赖SPA通过EventSource消费实时流式文本，'
        '集成修改意图检测、对话历史管理和实时增量预览。'
        '\n2. 拟解决的关键问题：'
        '（1）多源异构文件的信息统一抽取问题——如何设计统一管线并将结果归一化为结构化语义表示。'
        '（2）长文生成的幻觉控制问题——四层防御如何有效协同，在不增加过多token开销的前提下最大化拦截率。'
        '（3）推理类问题的检索增强问题——HyDE生成质量和思维链推理准确性的保证。'
        '（4）人机协作的细粒度控制问题——智能体自主执行与用户手动介入之间的平衡。'
        '\n3. 预期成果：'
        '（1）系统原型：完成可部署运行的AI文档生成智能体系统，包含Web前端和FastAPI后端，代码开源。'
        '（2）核心指标：支持5种以上输入格式，生成报告≥3章且单章正文≥200字，流式首字延迟<2秒，'
        '幻觉率（编造数字/机构名/人名）<10%（经四层防御后），并行文件抽取吞吐量4文件/批次。'
        '（3）技术文档：形成完整的技术架构文档，涵盖功能模块、系统架构、数据库设计和关键算法流程。'
        '（4）学术产出：就"长文生成中的幻觉抑制"方向撰写学术论文1篇。'
        '（5）可复用模块：rag/子包作为独立RAG知识库引擎可解耦复用。'
    ),
    '（七）项目研究进度安排': (
        '项目总周期为20周，分为七个阶段有序推进：'
        '\n第一阶段（第1-2周）：基础调研与技术选型。调研LangChain/LangGraph等LLM应用框架、HyDE与思维链推理等RAG前沿方案、'
        'PyMuPDF与RapidOCR等多格式文件处理库，形成技术选型报告并完成开发环境搭建。'
        '\n第二阶段（第3-5周）：多格式抽取引擎开发。依次实现PDF（含扫描件OCR）、Word、PPT、Excel和图片五类文件的解析与信息提取，'
        '搭建基于ThreadPoolExecutor的并行抽取架构，统一输出为六字段结构化结果，并通过单元测试验证正确性。'
        '\n第三阶段（第6-9周）：智能体编排核心开发。基于LangGraph实现ReAct智能体，定义七个工具并调试协同调用逻辑，'
        '完成大纲生成模块（含标题层级校验和自动重试）和正文生成模块（串行H2+并行H3策略），'
        '实现每个H2完成后的实时增量写入功能。'
        '\n第四阶段（第10-12周）：RAG推理与幻觉防御。实现HyDE假设文档嵌入、五步思维链Prompt、复杂问题自动分解与综合回答，'
        '将四层幻觉防御机制逐一落地，形成可独立运行的RAG知识库子包。'
        '\n第五阶段（第13-15周）：Web服务与前端交互界面。开发FastAPI全部REST路由（含SSE流式对话端点），'
        '实现前端SPA的文件上传拖拽、任务状态实时轮询、对话面板流式消费、Markdown实时渲染预览和知识库管理等交互模块。'
        '\n第六阶段（第16-17周）：全系统集成测试与性能优化。选取多种真实业务场景的文档样本执行端到端生成测试，'
        '调优并行抽取和并行生成的线程池参数，评估幻觉率并针对性调整防御判定阈值。'
        '\n第七阶段（第18-20周）：学术论文撰写与项目结题准备。以长文生成幻觉抑制为研究方向撰写论文并投稿，'
        '完善系统部署文档和用户操作手册，整理项目总结报告和答辩演示材料。'
    ),
    '（八）已有基础': (
        '1. 与本项目有关的研究积累和已取得的成绩。'
        '本项目已完成技术原型开发（agent-file-create v0.1.0），具备以下基础——'
        '多格式文件处理：已实现PDF（含OCR扫描件）、Word、PPT、Excel、图片5种格式的抽取器，并行处理架构就绪。'
        'LangGraph智能体：已实现7个工具的ReAct智能体，包含材料抽取→大纲生成→正文生成→模板渲染的完整生成管线，'
        '支持断点恢复。RAG知识库：已实现SQLite/PostgreSQL双后端向量存储、7级递归分块策略、三路混合检索+RRF融合、'
        'Cross-encoder/LLM双通道重排序、HyDE假设文档嵌入和思维链回答。长文生成增强：已实现四层幻觉防御体系、'
        'LLM语义摘要逻辑衔接、终审一致性检查、正则事实交叉比对。Web交互系统：已实现FastAPI+SSE流式对话、'
        '对话历史摘要与压缩、修改意图检测（40+关键词）、单章节定向重生成、实时增量预览。数据持久化：已实现双数据库'
        '架构（应用数据库5表+知识库数据库2表），支持SQLite和PostgreSQL。技术文档：已完成完整的技术架构设计文档'
        '和4篇专题技术文档（LangChain集成/智能体记忆/长文生成/推理性RAG），总计约2万字。'
        '代码规模约36个Python源文件，核心功能代码约8000+行。'
        '\n2. 已具备的条件及尚缺少的条件。'
        '已具备的条件：拥有可运行的技术原型agent-file-create v0.1.0，核心生成管线从材料上传到报告输出全链路贯通；'
        '开发环境配备Windows 11开发主机并启用WSL2 Linux子系统，项目运行于Python 3.12环境中；'
        '通过DeepSeek API可访问高性能模型，同时在本地通过Ollama部署了bge-m3嵌入模型、minicpm-v:8b视觉模型等；'
        '本地部署了PostgreSQL 15并安装pgvector扩展；已有一定数量的多格式测试文件样本；项目已产出完整的技术架构设计文档。'
        '\n尚缺少的条件及解决方法：GPU服务器资源——当前通过DeepSeek云端API调用可绕过此瓶颈，后续计划向学校实验室申请GPU资源；'
        '大规模测试数据——计划从arXiv和PMC等学术公开数据库及政府公开报告中收集素材构建测试集；'
        '中文长文本幻觉率评估标准——将在调研FACTOID、WikiBio等方法基础上设计结合人工标注和自动化规则检查的评估方案；'
        '真实用户反馈——计划在项目后期邀请5-10名同学进行可用性测试，基于反馈对交互流程进行迭代优化。'
    ),
    '（九）经费预算': (
        '项目申请总经费20000元，具体预算如下：'
        '\n1. DeepSeek API调用费用：5000元。用于大语言模型（deepseek-v4-pro、deepseek-v4-flash等）的文本生成、'
        '信息抽取和推理问答等API调用，覆盖开发测试与系统运行期间的token消耗。'
        '\n2. 云GPU服务器租赁：6000元。用于本地大模型（bge-m3嵌入模型、minicpm-v:8b视觉模型等）的低延迟推理部署，'
        '以及大规模测试数据的批量处理。'
        '\n3. 文献资料与数据库费用：3000元。用于购买技术书籍、学术论文数据库访问权限（如IEEE Xplore、ACM Digital Library）、'
        '以及获取开源许可的专业数据资源。'
        '\n4. 软件工具与开发环境：2000元。用于购买必要的开发工具（如PyCharm Professional许可证）、'
        '测试环境搭建（云服务器临时实例）和版本管理服务。'
        '\n5. 学术交流与论文发表：3000元。用于参加国内学术会议（如CCF中国软件大会、全国知识图谱与语义计算大会等）的'
        '注册费与差旅费，以及期刊论文的版面费。'
        '\n6. 项目管理与杂项：1000元。用于项目日常运行开支，包括打印耗材、办公用品、团队成员通讯补贴等。'
    ),
}

for section_title, content in sections_text.items():
    add_paragraph(section_title, bold=True, size=12, font_name='宋体', first_line_indent_cm=0.74, spacing_after=3)
    # Split content by the \n markers and add as separate paragraphs
    paragraphs = content.split('\n')
    for para_text in paragraphs:
        para_text = para_text.strip()
        if para_text and not para_text.startswith('（'):
            add_paragraph(para_text, size=12, font_name='宋体', first_line_indent_cm=0.74, spacing_after=3)
        elif para_text:
            add_paragraph(para_text, size=12, font_name='宋体', first_line_indent_cm=0.74, spacing_after=3)
    add_paragraph('', size=6, spacing_after=0)

add_paragraph('')

# ══════════════════════════════════════════════════════════════════
# 四、指导教师意见
# ══════════════════════════════════════════════════════════════════
add_paragraph('四、指导教师意见', bold=True, size=14, font_name='黑体', spacing_after=6)

opinion_table = doc.add_table(rows=1, cols=1)
opinion_table.alignment = WD_TABLE_ALIGNMENT.CENTER
opinion_table.style = 'Table Grid'
opinion_cell = opinion_table.cell(0, 0)
opinion_cell.width = Cm(14)
tr = opinion_table.rows[0]._tr
trPr = tr.get_or_add_trPr()
trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="2000" w:hRule="atLeast"/>')
trPr.append(trHeight)

add_cell_paragraphs(opinion_cell, [
    '',
    '',
    '本项目研究目标明确，技术方案可行，已具备良好的研究基础和技术原型。',
    '同意指导。',
    '',
    '指导教师签名：                                ',
    '日期：       年    月    日',
], size=12, font_name='宋体')

add_paragraph('')

# ══════════════════════════════════════════════════════════════════
# 五、院系推荐意见
# ══════════════════════════════════════════════════════════════════
add_paragraph('五、院系推荐意见', bold=True, size=14, font_name='黑体', spacing_after=6)

dept_table = doc.add_table(rows=1, cols=1)
dept_table.alignment = WD_TABLE_ALIGNMENT.CENTER
dept_table.style = 'Table Grid'
dept_cell = dept_table.cell(0, 0)
dept_cell.width = Cm(14)
tr = dept_table.rows[0]._tr
trPr = tr.get_or_add_trPr()
trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="1500" w:hRule="atLeast"/>')
trPr.append(trHeight)

add_cell_paragraphs(dept_cell, [
    '',
    '',
    '负责人签名：                                  ',
    '（盖院系章）                                   ',
    '日期：       年    月    日',
], size=12, font_name='宋体')

add_paragraph('')

# ══════════════════════════════════════════════════════════════════
# 六、学校推荐意见
# ══════════════════════════════════════════════════════════════════
add_paragraph('六、学校推荐意见', bold=True, size=14, font_name='黑体', spacing_after=6)

school_table = doc.add_table(rows=1, cols=1)
school_table.alignment = WD_TABLE_ALIGNMENT.CENTER
school_table.style = 'Table Grid'
school_cell = school_table.cell(0, 0)
school_cell.width = Cm(14)
tr = school_table.rows[0]._tr
trPr = tr.get_or_add_trPr()
trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="1500" w:hRule="atLeast"/>')
trPr.append(trHeight)

add_cell_paragraphs(school_cell, [
    '',
    '',
    '负责人签名：                                  ',
    '（盖学校章）                                   ',
    '日期：       年    月    日',
], size=12, font_name='宋体')

# ── Save ──
output_path = 'docs/创新训练项目申请书-叶永青.docx'
doc.save(output_path)
print(f'Application saved to: {output_path}')
