"""Generate the architecture design document (.docx)."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Global style settings ───────────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "微软雅黑"
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "微软雅黑"
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# ── Helper functions ────────────────────────────────────────────────────────
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text, bold=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val or "")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacer
    return table

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


# ╔════════════════════════════════════════════════════════════════════════════╗
# ║  封面                                                                     ║
# ╚════════════════════════════════════════════════════════════════════════════╝
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("AI 文档生成智能体\n架构设计文档")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("agent-file-create 项目技术设计与实现方案")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("版本：0.1.0  |  日期：2026-05-17  |  框架：LangChain + LangGraph + FastAPI").font.size = Pt(10)

doc.add_page_break()

# ╔════════════════════════════════════════════════════════════════════════════╗
# ║  目录占位                                                                ║
# ╚════════════════════════════════════════════════════════════════════════════╝
add_heading("目录", 1)
add_para("1. 功能模块架构设计\n2. 系统架构设计\n3. 数据库表设计")
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 第一章：功能模块架构
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("一、功能模块架构设计", 1)

add_heading("1.1 总体模块划分", 2)
add_para("agent-file-create 采用领域驱动的模块化设计，按功能职责将系统划分为 8 个核心模块，"
         "各模块遵循单向依赖原则：底层模块不依赖上层模块。")

add_heading("1.2 模块架构图", 2)
add_code("""┌─────────────────────────────────────────────────────────────────────┐
│                        Web 层  (agent_file_create/web/)              │
│   FastAPI REST API  │  SSE 流式对话  │  静态文件服务  │  任务调度入口  │
└───────────────────────────────┬─────────────────────────────────────┘
                                │
        ┌───────────────────────┼───────────────────────┐
        ▼                       ▼                       ▼
┌───────────────┐   ┌───────────────────┐   ┌───────────────────────┐
│  Chat 对话层   │   │  Agent 智能体层    │   │  Document 文档生成层   │
│  handler.py   │   │  document_agent   │   │  extractor.py         │
│  history.py   │   │  (LangGraph)      │   │  outline_generator.py │
│  prompts.py   │   │  7 Tools + ReAct  │   │  content_generator.py │
│               │   │  MemorySaver      │   │  template_renderer.py │
└───────┬───────┘   └────────┬──────────┘   └───────────┬───────────┘
        │                    │                          │
        │                    ▼                          │
        │   ┌────────────────────────────┐               │
        │   │  Task 任务管理层            │               │
        │   │  manager.py                │               │
        │   │  状态机 │ 控制事件 │ 持久化 │               │
        │   └────────────┬───────────────┘               │
        │                │                               │
        ▼                ▼                               ▼
┌─────────────────────────────────────────────────────────────────────┐
│                      LLM 基础设施层                                   │
│  llm_factory.py (LRU缓存工厂)  │  llm_client.py (HTTP调用)          │
│  config.py (环境变量配置中心)   │  prompts.py (Prompt模板)           │
└─────────────────────────────────────────────────────────────────────┘
        │
        ▼
┌─────────────────────────────────────────────────────────────────────┐
│                      RAG 知识库层  (agent_file_create/rag/)          │
│  kb.py (编排器)  │  chunker.py (分块)  │  embedder.py (向量化)       │
│  store.py (SQLite/PG) │  reranker.py (重排序) │  retriever.py (检索) │
└─────────────────────────────────────────────────────────────────────┘""")

add_heading("1.3 模块职责详述", 2)

add_table(
    ["模块", "目录/文件", "职责", "核心类/函数"],
    [
        ["Web 服务层", "web/server.py",
         "REST API 路由注册、请求参数校验、SSE 流式响应封装、\n静态文件挂载、后台任务线程管理",
         "FastAPI app\n_run_task()\n_start_task_thread()\n_make_regenerate_fn()"],
        ["Chat 对话层", "chat/handler.py\nchat/history.py\nchat/prompts.py",
         "对话上下文构建（三级可信度）、系统指令解析（/pause 等20+命令）、\n对话历史自动摘要压缩、追问推荐生成、修改意图检测、\nLangChain 历史适配器",
         "ChatHandler\nTaskChatMessageHistory\nlobby_prompt / task_chat_prompt"],
        ["Agent 智能体层", "agent/document_agent.py",
         "LangGraph ReAct 智能体、7 个工具定义（抽取/评估/大纲/\n正文/渲染/询问/完成）、MemorySaver 检查点、\n人机交互轮次管理",
         "DocumentAgent\n_build_tools()\nrun()"],
        ["Document 文档生成层", "document/extractor.py\ndocument/outline_generator.py\ndocument/content_generator.py\ndocument/template_renderer.py",
         "多格式文件信息抽取（PDF/图片/Word/PPT/Excel）、\n大纲生成与结构校验、分章节正文生成（串行H2+并行H3）、\nLLM摘要逻辑链衔接、终审一致性检查、\n模板渲染（md/docx/pdf）",
         "extract_from_file()\ngenerate_outline()\ngenerate_full_content_parallel()\nregenerate_section()\nrender_template()"],
        ["Task 任务管理层", "task/manager.py",
         "任务生命周期管理（queued→processing→finished/failed/canceled）、\n状态持久化（status.json）、元数据管理（task_meta.json）、\npause/resume/cancel 线程控制事件、\n对话历史读写与摘要存储、澄清答案轮询等待",
         "TaskManager\nwrite_status() / read_status()\nwait_for_clarify()\nget_control_events()"],
        ["LLM 基础设施层", "llm_factory.py\nllm_client.py\nconfig.py\nprompts.py",
         "统一 LLM 工厂函数（@lru_cache 缓存）、\nChatOpenAI / ChatOllama 双后端、\n环境变量配置中心（20+ 配置项）、\n提取/规划 Prompt 模板构建",
         "get_chat_model()\ncall_llm()\nbuild_extract_prompt()\nbuild_planner_prompt()"],
        ["RAG 知识库层", "rag/kb.py\nrag/chunker.py\nrag/embedder.py\nrag/store.py\nrag/reranker.py\nrag/retriever.py\nrag/embeddings_lc.py",
         "知识库编排（摄入/搜索/回答）、7级递归文本分块、\n双后端向量嵌入（Ollama/OpenAI）、\nSQLite/PostgreSQL+pgvector 向量存储、\nCross-encoder + LLM 重排序、\nLangChain BaseRetriever 封装、\nHyDE 假设文档扩展 + 思维链推理回答",
         "KnowledgeBase\nchunk_text()\nembed_texts()\nSQLiteVectorStore\nPostgresVectorStore\nrerank()\nKnowledgeBaseRetriever"],
        ["数据持久化层", "db_service.py",
         "应用数据库（SQLite/PostgreSQL）、5张业务表、\n任务 CRUD、大纲/正文/渲染输出持久化",
         "init_db()\ncreate_task()\nsave_outline()\nsave_content()"],
    ],
    col_widths=[2.5, 3.5, 5.5, 4.5]
)

add_heading("1.4 模块依赖规则", 2)
add_para("依赖方向（箭头从依赖方指向被依赖方）：", bold=True)
add_code("""web/server.py ──────► chat/handler.py ──────► rag/retriever.py ──────► rag/kb.py ──────► rag/store.py
       │                      │                                               │
       │                      ├──► rag/kb.py（lobby 直接检索）                  ├──► rag/chunker.py
       │                      ├──► llm_factory.py                              ├──► rag/embedder.py
       │                      └──► task/manager.py                             └──► rag/reranker.py
       │
       ├──► task/manager.py
       ├──► document/content_generator.py ────► llm_client.py
       └──► agent/document_agent.py ──────────► document/ (extractor/outline/content/template)""")
add_para("关键约束：")
add_para("• rag/ 是独立子包，不依赖 chat/、document/、web/ 模块", size=9)
add_para("• chat/ 依赖 rag/ 和 task/，但不依赖 document/", size=9)
add_para("• web/ 是所有模块的组装层（最高层），负责编排所有模块", size=9)
add_para("• content_generator.py 保留 llm_client.py 调用（历史兼容），其他模块统一使用 llm_factory.get_chat_model()", size=9)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 第二章：系统架构
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("二、系统架构设计", 1)

add_heading("2.1 整体架构概览", 2)
add_para("系统采用 B/S 架构 + 异步任务处理模式。前端为单页面应用（SPA），后端为 FastAPI REST 服务，"
         "文档生成任务在独立线程中异步执行，前端通过轮询获取实时状态。")

add_heading("2.2 系统架构图", 2)
add_code("""┌──────────────────────────────────────────────────────────────────────────┐
│                           Frontend (SPA)                                    │
│  html/index.html  │  app.js  │  styles.css                                  │
│  ┌──────────────┬──────────────┬──────────────┬──────────────┬──────────┐  │
│  │ 文件上传/拖拽  │ 任务状态轮询  │ 对话面板(SSE) │ 预览(MD渲染) │ KB管理    │  │
│  └──────────────┴──────────────┴──────────────┴──────────────┴──────────┘  │
└──────────────────────────────────┬───────────────────────────────────────┘
                                   │ HTTP/SSE
                                   ▼
┌──────────────────────────────────────────────────────────────────────────┐
│                        Backend (FastAPI)                                    │
│                                                                             │
│  ┌─────────────────────────────────────────────────────────────────────┐   │
│  │                        Web 路由层  (server.py)                        │   │
│  │                                                                       │   │
│  │  POST /api/upload     ─── 上传文件 + 启动生成任务                      │   │
│  │  POST /api/append     ─── 追加文件到现有任务 + 重新生成                 │   │
│  │  GET  /api/status     ─── 任务状态查询（状态/进度/阶段/文件）           │   │
│  │  GET  /api/tasks      ─── 任务列表 + 单任务详情                        │   │
│  │  POST /api/clarify    ─── 提交澄清问题答案                             │   │
│  │  POST /api/chat       ─── 对话（JSON 同步响应）                        │   │
│  │  POST /api/chat/stream─── 对话（SSE 流式响应）                         │   │
│  │  GET  /api/chat/history── 对话历史 + 摘要                              │   │
│  │  POST /api/kb/upload  ─── 知识库文件摄入                               │   │
│  │  POST /api/kb/query   ─── 知识库查询（RAG）                            │   │
│  │  GET  /api/kb/list    ─── 知识库列表                                   │   │
│  │  GET  /api/kb/docs    ─── 知识库文档列表                               │   │
│  │  POST /api/kb/delete  ─── 删除知识库/文档                              │   │
│  │  GET  /api/kb/stats   ─── 知识库统计                                   │   │
│  └─────────────────────────────────────────────────────────────────────┘   │
│                                   │                                         │
│         ┌─────────────────────────┼─────────────────────────┐               │
│         ▼                         ▼                         ▼               │
│  ┌──────────────┐    ┌──────────────────────┐    ┌──────────────────┐      │
│  │  ChatHandler │    │   DocumentAgent      │    │  TaskManager     │      │
│  │  (对话编排)   │    │   (LangGraph ReAct)   │    │  (任务生命周期)   │      │
│  │              │    │                      │    │                  │      │
│  │ • 上下文构建  │    │  7 tools:             │    │ • 状态机         │      │
│  │ • 指令解析    │    │  extract_files       │    │ • 控制事件       │      │
│  │ • 历史摘要    │    │  assess_material     │    │ • 对话持久化     │      │
│  │ • 修改意图    │    │  generate_outline    │    │ • 澄清轮询       │      │
│  │ • 追问推荐    │    │  generate_content    │    │ • 文件管理       │      │
│  │ • 流式回复    │    │  render_templates    │    │                  │      │
│  └──────┬───────┘    │  ask_user            │    └────────┬─────────┘      │
│         │            │  finish               │            │                │
│         │            └──────────┬───────────┘            │                │
│         │                       │                        │                │
│         └───────────────────────┼────────────────────────┘                │
│                                 ▼                                          │
│  ┌──────────────────────────────────────────────────────────────────────┐  │
│  │                     Document Pipeline                                 │  │
│  │                                                                       │  │
│  │  extractor.py ──► outline_generator.py ──► content_generator.py       │  │
│  │  (信息抽取)        (大纲生成)               (正文生成)                  │  │
│  │       │                   │                      │                    │  │
│  │       │   ┌───────────────┘                      │                    │  │
│  │       │   │                                      │                    │  │
│  │       ▼   ▼                                      ▼                    │  │
│  │  document_service.py (编排器) ──────► template_renderer.py            │  │
│  │                                         (模板渲染: md/docx/pdf)       │  │
│  └──────────────────────────────────────────────────────────────────────┘  │
│                                                                             │
│  ┌──────────────────────────────────────────────────────────────────────┐  │
│  │                        RAG Pipeline                                   │  │
│  │                                                                       │  │
│  │  文档 ──► chunker.py ──► embedder.py ──► store.py                      │  │
│  │            (7级递归分块)   (向量嵌入)       (SQLite/PG+pgvector)       │  │
│  │                                                                       │  │
│  │  查询 ──► 改写/ HyDE ──► 向量+词汇+BM25 ──► RRF融合 ──► 重排序        │  │
│  │            扩展            三路召回          (k=60)    cross-enc/LLM   │  │
│  └──────────────────────────────────────────────────────────────────────┘  │
└──────────────────────────────────────────────────────────────────────────┘""")

add_heading("2.3 技术栈", 2)
add_table(
    ["层次", "技术选型", "说明"],
    [
        ["前端", "原生 HTML5 + CSS3 + JavaScript (ES6+)", "零依赖 SPA，SSE 流式消费，Markdown 渲染，拖拽上传"],
        ["Web 框架", "FastAPI + Uvicorn", "异步 REST API，原生 SSE 支持，自动 OpenAPI 文档生成"],
        ["AI 框架", "LangChain Core + LangGraph", "ChatPromptTemplate 模板化、LCEL 链式组合、RunnableWithMessageHistory 历史管理、create_react_agent 智能体构建"],
        ["LLM 后端", "DeepSeek v4 (OpenAI 兼容 API)\nOllama 本地模型", "PLANNER_MODEL: deepseek-v4-pro\nCONTENT_MODEL: deepseek-v4-flash\nEMBED_MODEL: bge-m3\nVISION_MODEL: minicpm-v:8b"],
        ["向量嵌入", "BGE-M3 (via Ollama/OpenAI)", "1024维向量，支持中英双语检索"],
        ["向量存储", "SQLite (本地) / PostgreSQL + pgvector (生产)", "支持 HNSW 索引（m=16, ef=64）或 IVFFlat 索引"],
        ["重排序", "BAAI/bge-reranker-v2-m3 + LLM 回退", "Cross-encoder 精确重排序，LLM 回退保障可用性"],
        ["OCR", "RapidOCR + PyMuPDF", "图片文字提取 + PDF 页面渲染识别"],
        ["文档处理", "PyMuPDF / python-docx / python-pptx / openpyxl", "多格式文件文本与结构提取"],
        ["模板渲染", "docxtpl / reportlab / jinja2", "docx 模板变量替换 + PDF 画布绘制 + Markdown 文本替换"],
        ["数据持久化", "SQLite / PostgreSQL", "应用元数据 + 知识库向量存储，双数据库架构"],
    ],
    col_widths=[2, 4, 10]
)

add_heading("2.4 数据流", 2)
add_heading("2.4.1 文档生成主流程", 3)
add_para("用户上传文件到最终输出文档的完整数据流：", size=9)
add_code("""用户文件 + prompt
    │
    ▼
[1] 并行信息抽取 (ThreadPoolExecutor)
    │  每个文件：图片→OCR+Vision / PDF→PyMuPDF+OCR / Excel→pandas+LLM
    │  Word/PPT→结构化解析 / 文本→LLM提取
    │  输出：[{title, keywords, summary, key_points, data, conclusion}]
    │
    ▼
[2] A/B 质量评估（可选）
    │  对比预处理前后的提取质量，择优使用
    │
    ▼
[3] 用户澄清（材料不足时触发）
    │  Agent 生成澄清问题 → write_status("need_user") → 轮询等待用户回答
    │
    ▼
[4] 大纲生成 (LLM PLANNER_MODEL)
    │  输入: analysis_results + user_prompt
    │  输出: Markdown 大纲 (# Title \n ## Section \n ### Sub-section ...)
    │  校验: 标题层级合法性 + 最少章节数检查，最多 3 次自动重试
    │
    ▼
[5] 正文生成 (LLM CONTENT_MODEL)
    │  Phase 1: H2 章节串行生成（保证主线逻辑）
    │          每个章节用前序章节的 LLM 摘要作为上下文
    │          每完成一个 H2 → 增量写入 content.md（前端实时预览）
    │  Phase 2: H3+ 子节并行生成 (ThreadPoolExecutor)
    │          共享父 H2 章节的滚动摘要
    │  终审: LLM 语义一致性检查 + 正则事实交叉比对
    │
    ▼
[6] 模板渲染
    │  .md  → {{variable}} 文本替换
    │  .docx → docxtpl 模板引擎渲染
    │  .pdf → reportlab 画布绘制
    │
    ▼
[7] 持久化 + 通知
    写入 content.md / outline.md / 渲染输出 → DB 记录 → status="finished\"""")
add_para("")

add_heading("2.4.2 对话交互流程", 3)
add_code("""用户消息 + task_id
    │
    ▼
[1] _is_trivial_message() ─── 是 ──► 返回简短预设回复（问候/社交）
    │  否
    ▼
[2] _parse_chat_action() ─── 是命令 ──► _handle_chat_action() → 直接返回结果
    │  否                          (/pause /status /regen /kb 等)
    ▼
[3] status=="need_user" ─── 是 ──► 验证澄清答案 → write clarify_answers → 返回
    │  否
    ▼
[4] _build_context() 构建上下文
    │  • 加载 outline + content → 检索相关片段 (可信度：高)
    │  • KB 知识库检索 → 匹配文档片段 (可信度：中)
    │  • 对话摘要加载 (可信度：低)
    │  • Lobby 模式无任务：直接 KB 检索或通用问答
    │
    ▼
[5] LCEL Chain 调用 (lobby_chain 或 task_chain)
    │  RunnableWithMessageHistory 自动注入/保存历史
    │
    ▼
[6] 后处理
    • 追问推荐生成（异步线程）
    • 修改意图检测 → 更新 user_prompt → 建议 /regen
    • 历史摘要触发检查（>16条 或 >2000 token）""")

add_heading("2.4.3 RAG 检索流程", 3)
add_code("""用户查询
    │
    ├── [改写] 短查询(<10字)或中长查询 → LLM 关键词扩展
    ├── [HyDE] 复杂推理问题 → LLM 生成假设回答 → 用假设回答的 embedding 检索
    │
    ▼
[三路并行召回]
    ├── 向量检索: cosine similarity (SQLite) / pgvector <=> (PG)
    ├── 词汇检索: ILIKE/LIKE %term% (基于 BM25 词项匹配)
    └── BM25 重打分: k1=1.2, b=0.75, IDF 平滑
    │
    ▼
[RRF 融合] k=60, score = 1/(60+vec_rank) + 1/(60+bm25_rank) + 1/(60+lex_rank)
    │
    ▼
[重排序]
    ├── 首选: FlagReranker (BGE-reranker-v2-m3) cross-encoder
    └── 回退: LLM 逐条打分排序
    │
    ▼
[后处理]
    • 每文档最多选 2 块（diversity 控制）
    • 相邻块合并为 segment
    • 上下文窗口扩展（context_window=2 拉取前后相邻块）
    • 总上下文上限 5200 字符
    │
    ▼
[回答生成]
    ├── 简单问题: _ANSWER_PROMPT → 直接检索 + 简洁回答
    ├── 复杂推理: _ANSWER_COT_PROMPT → 5步思维链推理
    └── 多角度: 问题分解 → 子问题独立检索 → 综合回答""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 第三章：数据库表设计
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("三、数据库表设计", 1)

add_para("系统使用双数据库架构：应用数据库（存储任务与文档元数据）和知识库数据库（存储向量与分块）。"
         "两者均支持 SQLite（轻量部署）和 PostgreSQL（生产环境）。")

add_heading("3.1 应用数据库（db_service.py）", 2)
add_para("数据库路径：SQLite 默认 <项目根>/db/agent_file_create.db；PostgreSQL 通过 DB_URL 环境变量配置。", size=9)

add_heading("3.1.1 表：document_tasks（文档任务）", 3)
add_table(
    ["字段", "类型", "约束", "说明"],
    [
        ["id", "TEXT", "PRIMARY KEY", "任务唯一标识（8-36位UUID/hex）"],
        ["title", "TEXT", "NOT NULL DEFAULT ''", "报告标题（生成后回填）"],
        ["document_type", "TEXT", "NOT NULL DEFAULT ''", "文档类型：report / analysis / summary"],
        ["user_prompt", "TEXT", "NOT NULL DEFAULT ''", "用户原始需求描述文本"],
        ["status", "TEXT", "NOT NULL DEFAULT ''",
         "任务状态：queued / processing / paused / need_user / finished / failed / canceled"],
        ["output_dir", "TEXT", "NOT NULL DEFAULT ''", "输出目录绝对路径"],
        ["meta_json", "TEXT", "NOT NULL DEFAULT '{}'", "扩展元数据 JSON：template_mode / active_kb / ab_eval 等"],
        ["created_at", "REAL / DOUBLE", "NOT NULL", "创建时间戳（Unix epoch seconds）"],
        ["updated_at", "REAL / DOUBLE", "NOT NULL", "最后更新时间戳"],
    ],
    col_widths=[2.5, 2.5, 4, 7]
)

add_heading("3.1.2 表：document_outlines（文档大纲）", 3)
add_table(
    ["字段", "类型", "约束", "说明"],
    [
        ["id", "TEXT", "PRIMARY KEY", "大纲唯一标识（uuid.hex）"],
        ["task_id", "TEXT", "NOT NULL", "关联 document_tasks.id"],
        ["outline_markdown", "TEXT", "NOT NULL DEFAULT ''", "完整 Markdown 大纲原文"],
        ["outline_tree_json", "TEXT", "NOT NULL DEFAULT '[]'", "大纲章节树结构 JSON"],
        ["created_at", "REAL / DOUBLE", "NOT NULL", "创建时间戳"],
    ],
    col_widths=[2.5, 2.5, 4, 7]
)

add_heading("3.1.3 表：outline_sections（大纲章节）", 3)
add_table(
    ["字段", "类型", "约束", "说明"],
    [
        ["id", "TEXT", "PRIMARY KEY", "章节唯一标识"],
        ["outline_id", "TEXT", "NOT NULL", "关联 document_outlines.id"],
        ["task_id", "TEXT", "NOT NULL", "关联 document_tasks.id（冗余，加速查询）"],
        ["level", "INTEGER", "NOT NULL DEFAULT 0", "层级：1=H1标题, 2=H2章节, 3=H3子节, ..."],
        ["title", "TEXT", "NOT NULL DEFAULT ''", "章节标题文本"],
        ["parent_title", "TEXT", "NOT NULL DEFAULT ''", "父章节标题（H1的父标题为空）"],
        ["order_index", "INTEGER", "NOT NULL DEFAULT 0", "同级排序序号"],
    ],
    col_widths=[2.5, 2.5, 4, 7]
)

add_heading("3.1.4 表：document_contents（文档正文）", 3)
add_table(
    ["字段", "类型", "约束", "说明"],
    [
        ["id", "TEXT", "PRIMARY KEY", "正文唯一标识"],
        ["task_id", "TEXT", "NOT NULL", "关联 document_tasks.id"],
        ["markdown_content", "TEXT", "NOT NULL DEFAULT ''", "完整 Markdown 正文"],
        ["meta_json", "TEXT", "NOT NULL DEFAULT '{}'", "扩展元数据：outline_id / output_dir / template_dir"],
        ["created_at", "REAL / DOUBLE", "NOT NULL", "创建时间戳"],
    ],
    col_widths=[2.5, 2.5, 4, 7]
)

add_heading("3.1.5 表：rendered_outputs（渲染输出文件）", 3)
add_table(
    ["字段", "类型", "约束", "说明"],
    [
        ["id", "TEXT", "PRIMARY KEY", "输出记录唯一标识"],
        ["task_id", "TEXT", "NOT NULL", "关联 document_tasks.id"],
        ["file_path", "TEXT", "NOT NULL DEFAULT ''", "渲染输出文件的绝对路径"],
        ["created_at", "REAL / DOUBLE", "NOT NULL", "创建时间戳"],
    ],
    col_widths=[2.5, 2.5, 4, 7]
)

add_heading("3.1.6 应用数据库 ER 关系", 3)
add_code("""document_tasks (1) ────── (N) document_outlines ────── (N) outline_sections
      │
      └─────────────────── (N) document_contents
      │
      └─────────────────── (N) rendered_outputs

关系说明：
• 一个任务可以多次生成大纲（重新生成时追加新记录，旧记录保留）
• 一个大纲包含多个章节（outline_sections 展开存储，支持按标题/层级查询）
• 一次生成产出一份正文（document_contents.meta_json.outline_id 关联大纲版本）
• 一次渲染可能产出多个格式文件（.md / .docx / .pdf）""")

add_heading("3.2 知识库数据库（store.py）", 2)
add_para("数据库路径：SQLite 默认 <项目根>/db/rag.db；PostgreSQL 通过 KB_DB_URL 环境变量配置。", size=9)

add_heading("3.2.1 表：kb_docs（知识库文档）", 3)
add_table(
    ["字段", "类型 (SQLite)", "类型 (PostgreSQL)", "约束", "说明"],
    [
        ["id", "TEXT", "TEXT", "NOT NULL", "文档唯一标识（文件名/MD5）"],
        ["kb", "TEXT", "TEXT", "NOT NULL", "知识库名称（命名空间隔离）"],
        ["title", "TEXT", "TEXT", "NOT NULL DEFAULT ''", "文档标题"],
        ["source", "TEXT", "TEXT", "NOT NULL DEFAULT ''", "文档来源路径/URL"],
        ["doc_type", "—", "TEXT", "DEFAULT ''", "文档类型：pdf/docx/pptx/xlsx/txt/md（仅PG）"],
        ["meta_json / meta", "TEXT", "JSONB", "DEFAULT '{}'", "扩展元数据（SQLite用TEXT存JSON）"],
        ["updated_at", "REAL", "TIMESTAMPTZ", "NOT NULL", "最后更新时间戳"],
    ],
    col_widths=[2.5, 2, 2.5, 3, 6]
)
add_para("主键：SQLite → id；PostgreSQL → (kb, id) 联合主键", size=9)
add_para("索引：kb_docs(kb)", size=9)

add_heading("3.2.2 表：kb_chunks（文档分块）", 3)
add_table(
    ["字段", "类型 (SQLite)", "类型 (PostgreSQL)", "约束", "说明"],
    [
        ["id", "TEXT", "TEXT", "PRIMARY KEY", "分块唯一标识：{doc_id}:{chunk_index} 或 {doc_id}:parent:{idx}"],
        ["kb", "TEXT", "TEXT", "NOT NULL", "知识库名称"],
        ["doc_id", "TEXT", "TEXT", "NOT NULL", "关联 kb_docs.id"],
        ["chunk_index", "INTEGER", "INTEGER", "NOT NULL", "分块序号（0-based）"],
        ["section_path", "TEXT", "TEXT", "NOT NULL DEFAULT ''", "章节路径：如「第三章 / 3.1 定义」"],
        ["content", "TEXT", "TEXT", "NOT NULL DEFAULT ''", "分块文本内容"],
        ["embedding / embedding_json", "TEXT (JSON)", "vector", "NOT NULL", "向量嵌入（SQLite用JSON文本；PG用pgvector原生类型）"],
        ["source", "—", "TEXT", "DEFAULT ''", "源文件名（仅PG）"],
        ["title", "—", "TEXT", "DEFAULT ''", "文档标题（冗余，仅PG）"],
        ["doc_type", "—", "TEXT", "DEFAULT ''", "文档类型（冗余，仅PG）"],
        ["meta_json / meta", "TEXT", "JSONB", "DEFAULT '{}'", "扩展元数据"],
        ["created_at", "REAL", "TIMESTAMPTZ", "NOT NULL", "创建时间戳"],
    ],
    col_widths=[3, 2, 2.5, 2.5, 6]
)
add_para("索引：", bold=True, size=9)
add_para("• kb_chunks(kb) — 按知识库筛选", size=9)
add_para("• kb_chunks(kb, doc_id) — 按文档查询所有分块（用于 context_window 扩展和文档删除）", size=9)
add_para("• PostgreSQL 可选向量索引 — HNSW (m=16, ef_construction=64) 或 IVFFlat (lists=100)", size=9)
add_para("• embedding 维度：1024（BGE-M3），PostgreSQL 创建索引时需指定 vector_cosine_ops", size=9)

add_heading("3.2.3 知识库数据库 ER 关系", 3)
add_code("""kb_docs (1) ────── (N) kb_chunks
  (kb, id)              (kb, doc_id)

关系说明：
• 一个文档包含多个分块（通常 10-500 块，取决于文档长度）
• 分块通过 parent_chunk_id 维护父子关系（每4个子块绑定一个父块 ID）
• 父块本身也存储在 kb_chunks 表中（通过 chunk_id 区分：普通块="{doc_id}:{idx}"，父块="{doc_id}:parent:{p_idx}"）
• 查询时通过 context_window=2 拉取命中块的 ±2 相邻块
• 删除文档时级联删除其所有分块""")

add_heading("3.3 文件系统持久化（TaskManager）", 2)
add_para("除数据库外，系统还通过文件系统维护任务运行时数据，存储在 result/<task_id>/ 目录下：", size=9)
add_table(
    ["文件/目录", "格式", "读写方", "内容"],
    [
        ["status.json", "JSON", "TaskManager", "任务实时状态：{task_id, status, stage, message, total_files, done_files, clarify_questions, clarify_answers, ab_results, updated_at}"],
        ["task_meta.json", "JSON", "TaskManager", "任务持久元数据：{file_paths, user_prompt, ab_eval, template_mode, template_dir, active_kb, saved_templates}"],
        ["analysis_results.json", "JSON", "TaskManager", "文件提取结果缓存（避免重复抽取）"],
        ["chat_history.json", "JSON Array", "TaskManager", "对话历史：[{role, content}, ...]（最多50条，摘要后截断到最近8条）"],
        ["chat_summary.txt", "Plain Text", "TaskManager", "多轮对话摘要（最大2000字符，多轮摘要合并累积）"],
        ["outline.md", "Markdown", "document_service", "生成的大纲文件"],
        ["content.md", "Markdown", "content_generator", "生成的正文（生成中增量写入，生成后终审更新）"],
        ["uploads/", "Directory", "server.py", "用户上传的原始文件（PDF/Word/图片等）"],
        ["template/", "Directory", "server.py", "用户上传的模板文件（.md/.docx/.pdf）"],
        ["*_rendered.md/.docx/.pdf", "Mixed", "template_renderer", "模板渲染后的最终输出文件"],
    ],
    col_widths=[3.5, 2, 3, 7.5]
)

add_heading("3.4 数据流转与一致性", 2)
add_para("1. 任务创建时：server.py 写入 task_meta.json + status.json → db_service 插入 document_tasks 行", size=9)
add_para("2. 生成过程中：status.json 实时更新（每章节、每阶段），前端每 2 秒轮询 GET /api/status", size=9)
add_para("3. 生成完成后：content.md / outline.md 写入文件系统 → db_service 插入 document_outlines + document_contents + rendered_outputs → status 更新为 finished", size=9)
add_para("4. 数据修复：如果 DB 记录丢失，可通过文件系统重建（result/<task_id>/ 目录自包含）", size=9)
add_para("5. 知识库数据：kb_docs + kb_chunks 与文件系统无直接关联，独立生命周期管理", size=9)

# ── Save ─────────────────────────────────────────────────────────────────────
output_path = "d:/python/review/agent-file-create/docs/architecture_design.docx"
doc.save(output_path)
print(f"Saved to: {output_path}")
